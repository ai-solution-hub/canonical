"""Tests for template filling -- writing responses into Word documents.

Tests cover:
  - Programmatic template filling (empty cells, formatting, word limits)
  - Fixture-based tests using real .docx files (simple-template, complex-template)
  - Formatting preservation (CRITICAL: never cell.text = "value")
  - Word limit enforcement and truncation
  - Multi-paragraph response writing
  - Post-fill validation
  - Edge cases (missing indices, empty responses, column mismatch)
  - Track Changes safe opening via open_document_safe
"""

import io
import pytest
from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl
import os
import sys
import tempfile
from unittest.mock import patch, MagicMock

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from fill_template import (
    fill_template,
    fill_pdf_template,
    fill_xlsx_template,
    _enforce_word_limit,
    _copy_cell_formatting,
    _validate_completed_document,
    _pdf_field_names_by_sequence,
    _xlsx_sheet_by_table_index,
)
from analyse_template import analyse_template
from pypdf import PdfReader


# ── Fixtures ─────────────────────────────────────────────────────────────

FIXTURES_DIR = os.path.join(
    os.path.dirname(__file__), '..', '..', '__tests__', 'fixtures'
)
SIMPLE_FIXTURE = os.path.join(FIXTURES_DIR, 'simple-template.docx')
COMPLEX_FIXTURE = os.path.join(FIXTURES_DIR, 'complex-template.docx')


@pytest.fixture
def simple_analysis():
    """Analyse the simple fixture to get field mappings."""
    return analyse_template(SIMPLE_FIXTURE)


@pytest.fixture
def complex_analysis():
    """Analyse the complex fixture to get field mappings."""
    return analyse_template(COMPLEX_FIXTURE)


@pytest.fixture
def temp_output(tmp_path):
    """Provide a temporary output path for filled documents."""
    return str(tmp_path / "output.docx")


# ── Programmatic template helpers ────────────────────────────────────────

def _create_formatted_template(path: str) -> str:
    """Create a template with specific formatting for testing preservation.

    Returns the path to the created file.
    """
    doc = Document()
    table = doc.add_table(rows=4, cols=2)

    # Header row
    table.rows[0].cells[0].text = "Question"
    table.rows[0].cells[1].text = "Response"

    # Row 1: question with specific font
    q1_para = table.rows[1].cells[0].paragraphs[0]
    q1_run = q1_para.add_run("What is your approach to security?")
    q1_run.font.name = "Arial"
    q1_run.font.size = Pt(11)
    q1_run.font.bold = True
    # Response cell empty

    # Row 2: question with italic
    q2_para = table.rows[2].cells[0].paragraphs[0]
    q2_run = q2_para.add_run("Describe your experience")
    q2_run.font.name = "Arial"
    q2_run.font.size = Pt(11)
    q2_run.font.italic = True
    # Response cell empty

    # Row 3: question with content already
    q3_para = table.rows[3].cells[0].paragraphs[0]
    q3_run = q3_para.add_run("Already filled")
    q3_run.font.name = "Arial"
    q3_run.font.size = Pt(11)
    table.rows[3].cells[1].text = "Existing content"

    doc.save(path)
    return path


# ══════════════════════════════════════════════════════════════════════════
# PROGRAMMATIC TEMPLATE TESTS
# ══════════════════════════════════════════════════════════════════════════


class TestFillTemplate:
    """Test the main fill_template function with programmatic templates."""

    def test_fills_empty_cell(self):
        """Response text written into empty cell."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "We use ISO 27001", "word_limit": None}]
                )

                assert result["fields_filled"] == 1
                assert result["fields_failed"] == 0

                # Verify content was written
                doc = Document(output_path)
                cell_text = doc.tables[0].rows[1].cells[1].text.strip()
                assert "ISO 27001" in cell_text
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_preserves_non_filled_cells(self):
        """Cells not in the mapping list remain unchanged."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                # Only fill row 1, leave row 3 (which has content) alone
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "New response", "word_limit": None}]
                )

                doc = Document(output_path)
                # Row 3 should still have original content
                existing = doc.tables[0].rows[3].cells[1].text.strip()
                assert existing == "Existing content"
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_enforces_word_limit(self):
        """Long responses truncated to word limit."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            long_text = " ".join(["word"] * 100)

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": long_text, "word_limit": 10}]
                )

                assert result["fields_filled"] == 1
                assert len(result["truncated"]) == 1
                assert result["truncated"][0]["limit"] == 10
                assert result["truncated"][0]["original_words"] == 100

                # Verify truncated text
                doc = Document(output_path)
                written = doc.tables[0].rows[1].cells[1].text.strip()
                assert len(written.split()) <= 10
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_handles_missing_table(self):
        """Graceful failure when table_index is out of range."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 99, "row_index": 0, "col_index": 0,
                      "response_text": "Text", "word_limit": None}]
                )

                assert result["fields_failed"] == 1
                assert len(result["errors"]) == 1
                assert "out of range" in result["errors"][0]["error"]
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_handles_missing_row(self):
        """Graceful failure when row_index is out of range."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 99, "col_index": 0,
                      "response_text": "Text", "word_limit": None}]
                )

                assert result["fields_failed"] == 1
                assert len(result["errors"]) == 1
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_empty_response_skipped(self):
        """Empty response text results in skip, not fill."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "", "word_limit": None}]
                )

                assert result["fields_skipped"] == 1
                assert result["fields_filled"] == 0
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_document_unchanged_on_error(self):
        """Original template file is never modified."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            # Get original file contents
            with open(tmp_in.name, "rb") as f:
                original_bytes = f.read()

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "New content", "word_limit": None}]
                )

                # Verify original was not modified
                with open(tmp_in.name, "rb") as f:
                    after_bytes = f.read()
                assert original_bytes == after_bytes
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_multiple_fields_filled(self):
        """Multiple fields can be filled in one call."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [
                        {"table_index": 0, "row_index": 1, "col_index": 1,
                         "response_text": "Security response", "word_limit": None},
                        {"table_index": 0, "row_index": 2, "col_index": 1,
                         "response_text": "Experience response", "word_limit": None},
                    ]
                )

                assert result["fields_filled"] == 2
                assert result["fields_failed"] == 0
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)


# ══════════════════════════════════════════════════════════════════════════
# FORMATTING PRESERVATION TESTS
# ══════════════════════════════════════════════════════════════════════════


class TestFormattingPreservation:
    """Tests that fill_template preserves cell formatting.

    CRITICAL: fill_template must NEVER use cell.text = "value".
    It must use paragraph.clear() + paragraph.add_run() to preserve formatting.
    """

    def test_font_name_copied_from_question_cell(self):
        """Filled response cell inherits font name from the question cell."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "Test response", "word_limit": None}]
                )

                doc = Document(output_path)
                response_cell = doc.tables[0].rows[1].cells[1]
                runs = response_cell.paragraphs[0].runs
                assert len(runs) >= 1
                # Font name should be inherited from question cell (Arial)
                assert runs[0].font.name == "Arial"
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_font_size_copied_from_question_cell(self):
        """Filled response cell inherits font size from the question cell."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "Test response", "word_limit": None}]
                )

                doc = Document(output_path)
                response_cell = doc.tables[0].rows[1].cells[1]
                runs = response_cell.paragraphs[0].runs
                assert len(runs) >= 1
                assert runs[0].font.size == Pt(11)
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_bold_copied_from_question_cell(self):
        """Row 1 question is bold — response should inherit bold."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "Bold response", "word_limit": None}]
                )

                doc = Document(output_path)
                runs = doc.tables[0].rows[1].cells[1].paragraphs[0].runs
                assert len(runs) >= 1
                assert runs[0].font.bold is True
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_italic_copied_from_question_cell(self):
        """Row 2 question is italic — response should inherit italic."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 2, "col_index": 1,
                      "response_text": "Italic response", "word_limit": None}]
                )

                doc = Document(output_path)
                runs = doc.tables[0].rows[2].cells[1].paragraphs[0].runs
                assert len(runs) >= 1
                assert runs[0].font.italic is True
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_font_colour_preserved(self):
        """Font colour from reference cell is copied to response."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Question"
            table.rows[0].cells[1].text = "Response"

            q_para = table.rows[1].cells[0].paragraphs[0]
            q_run = q_para.add_run("Coloured question")
            q_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            doc.save(tmp_in.name)

            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "Coloured response", "word_limit": None}]
                )

                doc = Document(output_path)
                runs = doc.tables[0].rows[1].cells[1].paragraphs[0].runs
                assert len(runs) >= 1
                assert runs[0].font.color.rgb == RGBColor(0x00, 0x33, 0x66)
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_uses_run_api_not_cell_text(self):
        """Verify the response is written via runs, not cell.text assignment.

        If cell.text were used, the cell would have exactly one paragraph
        with text set directly (no runs with formatting). Using the run API
        creates a run object with formatting properties preserved.
        """
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "Written via run API", "word_limit": None}]
                )

                doc = Document(output_path)
                cell = doc.tables[0].rows[1].cells[1]
                # If the run API was used, the paragraph should have at least 1 run
                assert len(cell.paragraphs[0].runs) >= 1
                assert cell.paragraphs[0].runs[0].text == "Written via run API"
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)


# ══════════════════════════════════════════════════════════════════════════
# MULTI-PARAGRAPH RESPONSE TESTS
# ══════════════════════════════════════════════════════════════════════════


class TestMultiParagraphResponse:
    """Test that multi-paragraph responses (double newline separated) are handled."""

    def test_double_newline_creates_paragraphs(self):
        """Response with \\n\\n creates multiple paragraphs in correct order."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            multi_para = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": multi_para, "word_limit": None}]
                )

                assert result["fields_filled"] == 1

                doc = Document(output_path)
                cell = doc.tables[0].rows[1].cells[1]
                # Should have 3 paragraphs with content in correct order
                para_texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                assert len(para_texts) == 3
                assert para_texts[0] == "First paragraph."
                assert para_texts[1] == "Second paragraph."
                assert para_texts[2] == "Third paragraph."
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_single_newline_kept_in_same_paragraph(self):
        """Single newlines within a paragraph are not split."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            single_newline = "Line one\nLine two"

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": single_newline, "word_limit": None}]
                )

                doc = Document(output_path)
                cell = doc.tables[0].rows[1].cells[1]
                # Should be kept as one paragraph (split on \n\n, not \n)
                non_empty = [p for p in cell.paragraphs if p.text.strip()]
                assert len(non_empty) == 1
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)


# ══════════════════════════════════════════════════════════════════════════
# FIXTURE-BASED FILL TESTS — SIMPLE TEMPLATE
# ══════════════════════════════════════════════════════════════════════════


class TestFillSimpleFixture:
    """Fill the simple-template.docx fixture and verify results."""

    def test_fill_all_fields(self, simple_analysis, temp_output):
        """Fill all 5 fields and verify each one is written."""
        mappings = []
        for i, field in enumerate(simple_analysis["fields"]):
            mappings.append({
                "table_index": field["table_index"],
                "row_index": field["row_index"],
                "col_index": field["col_index"],
                "response_text": f"Response to field {i}",
                "word_limit": field.get("word_limit"),
            })

        result = fill_template(SIMPLE_FIXTURE, temp_output, mappings)

        assert result["fields_filled"] == 5
        assert result["fields_skipped"] == 0
        assert result["fields_failed"] == 0
        assert result["truncated"] == []
        assert result["errors"] == []

    def test_filled_fields_have_correct_text(self, simple_analysis, temp_output):
        """Each field contains the expected response text after filling."""
        mappings = [
            {
                "table_index": f["table_index"],
                "row_index": f["row_index"],
                "col_index": f["col_index"],
                "response_text": f"Answer for: {f['question_text'][:30]}",
                "word_limit": None,
            }
            for f in simple_analysis["fields"]
        ]

        fill_template(SIMPLE_FIXTURE, temp_output, mappings)

        doc = Document(temp_output)
        for field, mapping in zip(simple_analysis["fields"], mappings):
            cell = doc.tables[field["table_index"]].rows[field["row_index"]].cells[field["col_index"]]
            assert mapping["response_text"] in cell.text

    def test_output_is_valid_docx(self, simple_analysis, temp_output):
        """Output file can be opened as a valid .docx document."""
        mappings = [{
            "table_index": simple_analysis["fields"][0]["table_index"],
            "row_index": simple_analysis["fields"][0]["row_index"],
            "col_index": simple_analysis["fields"][0]["col_index"],
            "response_text": "Valid document test",
            "word_limit": None,
        }]

        fill_template(SIMPLE_FIXTURE, temp_output, mappings)

        # Should not raise any exception
        doc = Document(temp_output)
        assert len(doc.tables) == 2
        assert os.path.getsize(temp_output) > 0

    def test_partial_fill_leaves_other_fields_unchanged(self, simple_analysis, temp_output):
        """Filling only some fields leaves the rest as-is."""
        # Only fill the first field
        first_field = simple_analysis["fields"][0]
        mappings = [{
            "table_index": first_field["table_index"],
            "row_index": first_field["row_index"],
            "col_index": first_field["col_index"],
            "response_text": "Only this one",
            "word_limit": None,
        }]

        fill_template(SIMPLE_FIXTURE, temp_output, mappings)

        doc = Document(temp_output)
        # First field should be filled
        filled_cell = doc.tables[first_field["table_index"]].rows[first_field["row_index"]].cells[first_field["col_index"]]
        assert "Only this one" in filled_cell.text

        # Other placeholder fields should still have placeholder text or be empty
        for field in simple_analysis["fields"][1:]:
            cell = doc.tables[field["table_index"]].rows[field["row_index"]].cells[field["col_index"]]
            # Should NOT contain "Only this one"
            assert "Only this one" not in cell.text

    def test_post_fill_validation_passes(self, simple_analysis, temp_output):
        """Validation of filled document shows no structural warnings."""
        mappings = [{
            "table_index": f["table_index"],
            "row_index": f["row_index"],
            "col_index": f["col_index"],
            "response_text": "Response",
            "word_limit": None,
        } for f in simple_analysis["fields"]]

        fill_template(SIMPLE_FIXTURE, temp_output, mappings)

        warnings = _validate_completed_document(SIMPLE_FIXTURE, temp_output)
        assert len(warnings) == 0, f"Unexpected validation warnings: {warnings}"

    def test_table_and_row_counts_preserved(self, simple_analysis, temp_output):
        """Filling does not alter table count or row counts."""
        mappings = [{
            "table_index": f["table_index"],
            "row_index": f["row_index"],
            "col_index": f["col_index"],
            "response_text": "Response",
            "word_limit": None,
        } for f in simple_analysis["fields"]]

        fill_template(SIMPLE_FIXTURE, temp_output, mappings)

        original = Document(SIMPLE_FIXTURE)
        filled = Document(temp_output)

        assert len(original.tables) == len(filled.tables)
        for orig_t, fill_t in zip(original.tables, filled.tables):
            assert len(orig_t.rows) == len(fill_t.rows)


# ══════════════════════════════════════════════════════════════════════════
# FIXTURE-BASED FILL TESTS — COMPLEX TEMPLATE
# ══════════════════════════════════════════════════════════════════════════


class TestFillComplexFixture:
    """Fill the complex-template.docx fixture and verify results."""

    def test_fill_all_identified_fields(self, complex_analysis, temp_output):
        """Fill all 6 identified fields from the complex fixture."""
        mappings = []
        for i, field in enumerate(complex_analysis["fields"]):
            mappings.append({
                "table_index": field["table_index"],
                "row_index": field["row_index"],
                "col_index": field["col_index"],
                "response_text": f"Complex response {i}",
                "word_limit": field.get("word_limit"),
            })

        result = fill_template(COMPLEX_FIXTURE, temp_output, mappings)

        assert result["fields_filled"] == 6
        assert result["fields_failed"] == 0
        assert result["errors"] == []

    def test_word_limit_truncation_in_complex(self, complex_analysis, temp_output):
        """Fields with word limits are truncated when response exceeds limit."""
        # Find a field with a word limit
        limited_fields = [
            f for f in complex_analysis["fields"]
            if f.get("word_limit") is not None
        ]
        assert len(limited_fields) >= 1, "Expected at least one field with word limit"

        field = limited_fields[0]
        word_limit = field["word_limit"]
        # Generate a response that exceeds the word limit
        long_response = " ".join(["quality"] * (word_limit + 50))

        mappings = [{
            "table_index": field["table_index"],
            "row_index": field["row_index"],
            "col_index": field["col_index"],
            "response_text": long_response,
            "word_limit": word_limit,
        }]

        result = fill_template(COMPLEX_FIXTURE, temp_output, mappings)

        assert result["fields_filled"] == 1
        assert len(result["truncated"]) == 1
        assert result["truncated"][0]["limit"] == word_limit

        # Verify the written text respects the limit
        doc = Document(temp_output)
        cell = doc.tables[field["table_index"]].rows[field["row_index"]].cells[field["col_index"]]
        written_words = cell.text.strip().split()
        assert len(written_words) <= word_limit

    def test_three_column_table_fill(self, complex_analysis, temp_output):
        """Table 1 (Ref | Question | Response) — filling col 2 works correctly."""
        table1_fields = [
            f for f in complex_analysis["fields"]
            if f["table_index"] == 1
        ]
        assert len(table1_fields) >= 1

        mappings = [{
            "table_index": f["table_index"],
            "row_index": f["row_index"],
            "col_index": f["col_index"],
            "response_text": "Three-column table response",
            "word_limit": None,
        } for f in table1_fields]

        result = fill_template(COMPLEX_FIXTURE, temp_output, mappings)
        assert result["fields_filled"] == len(table1_fields)

        # Verify the response was written into the correct column
        doc = Document(temp_output)
        for field in table1_fields:
            cell = doc.tables[field["table_index"]].rows[field["row_index"]].cells[field["col_index"]]
            assert "Three-column table response" in cell.text

    def test_four_column_table_fill(self, complex_analysis, temp_output):
        """Table 2 (Ref | Question | Word Limit | Response) — filling col 3 works."""
        table2_fields = [
            f for f in complex_analysis["fields"]
            if f["table_index"] == 2
        ]
        assert len(table2_fields) >= 1

        mappings = [{
            "table_index": f["table_index"],
            "row_index": f["row_index"],
            "col_index": f["col_index"],
            "response_text": "Four-column response",
            "word_limit": None,
        } for f in table2_fields]

        result = fill_template(COMPLEX_FIXTURE, temp_output, mappings)
        assert result["fields_filled"] == len(table2_fields)

    def test_complex_post_fill_validation(self, complex_analysis, temp_output):
        """Validation passes for the complex fixture after filling."""
        mappings = [{
            "table_index": f["table_index"],
            "row_index": f["row_index"],
            "col_index": f["col_index"],
            "response_text": "Validation test response",
            "word_limit": None,
        } for f in complex_analysis["fields"]]

        fill_template(COMPLEX_FIXTURE, temp_output, mappings)

        warnings = _validate_completed_document(COMPLEX_FIXTURE, temp_output)
        assert len(warnings) == 0, f"Unexpected validation warnings: {warnings}"

    def test_complex_table_counts_preserved(self, complex_analysis, temp_output):
        """Filling does not alter the complex document's table/row structure."""
        mappings = [{
            "table_index": f["table_index"],
            "row_index": f["row_index"],
            "col_index": f["col_index"],
            "response_text": "Structure preservation test",
            "word_limit": None,
        } for f in complex_analysis["fields"]]

        fill_template(COMPLEX_FIXTURE, temp_output, mappings)

        original = Document(COMPLEX_FIXTURE)
        filled = Document(temp_output)

        assert len(original.tables) == len(filled.tables)
        for t_idx, (orig_t, fill_t) in enumerate(zip(original.tables, filled.tables)):
            assert len(orig_t.rows) == len(fill_t.rows), (
                f"Table {t_idx} row count changed: {len(orig_t.rows)} -> {len(fill_t.rows)}"
            )

    def test_formatting_preserved_in_complex_fixture(self, complex_analysis, temp_output):
        """Formatting (font properties) preserved in the complex fixture after filling."""
        # Fill the first field in table 0
        table0_fields = [f for f in complex_analysis["fields"] if f["table_index"] == 0]
        if not table0_fields:
            pytest.skip("No table 0 fields available")

        field = table0_fields[0]
        mappings = [{
            "table_index": field["table_index"],
            "row_index": field["row_index"],
            "col_index": field["col_index"],
            "response_text": "Formatted response",
            "word_limit": None,
        }]

        fill_template(COMPLEX_FIXTURE, temp_output, mappings)

        doc = Document(temp_output)
        cell = doc.tables[field["table_index"]].rows[field["row_index"]].cells[field["col_index"]]
        runs = cell.paragraphs[0].runs
        # Should have at least one run (written via run API, not cell.text)
        assert len(runs) >= 1
        assert runs[0].text == "Formatted response"


# ══════════════════════════════════════════════════════════════════════════
# MIXED RESULTS TESTS
# ══════════════════════════════════════════════════════════════════════════


class TestMixedResults:
    """Tests with a mix of filled, skipped, and failed fields."""

    def test_mixed_fill_skip_fail(self):
        """One valid fill, one empty skip, one out-of-range fail."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [
                        # Valid fill
                        {"table_index": 0, "row_index": 1, "col_index": 1,
                         "response_text": "Filled", "word_limit": None},
                        # Skip (empty response)
                        {"table_index": 0, "row_index": 2, "col_index": 1,
                         "response_text": "", "word_limit": None},
                        # Fail (out of range)
                        {"table_index": 99, "row_index": 0, "col_index": 0,
                         "response_text": "Bad index", "word_limit": None},
                    ]
                )

                assert result["fields_filled"] == 1
                assert result["fields_skipped"] == 1
                assert result["fields_failed"] == 1
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_whitespace_only_response_skipped(self):
        """Response containing only whitespace is treated as empty (skipped)."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "   \n\t  ", "word_limit": None}]
                )

                assert result["fields_skipped"] == 1
                assert result["fields_filled"] == 0
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_none_response_skipped(self):
        """None response_text is treated as empty (skipped)."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": None, "word_limit": None}]
                )

                assert result["fields_skipped"] == 1
                assert result["fields_filled"] == 0
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_column_index_out_of_range(self):
        """Graceful failure when col_index is out of range."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                result = fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 99,
                      "response_text": "Bad column", "word_limit": None}]
                )

                assert result["fields_failed"] == 1
                assert "out of range" in result["errors"][0]["error"]
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)


# ══════════════════════════════════════════════════════════════════════════
# WORD LIMIT HELPER TESTS
# ══════════════════════════════════════════════════════════════════════════


class TestEnforceWordLimit:
    """Test the word limit enforcement helper."""

    def test_no_limit(self):
        text, truncated = _enforce_word_limit("Hello world", None)
        assert text == "Hello world"
        assert truncated is False

    def test_within_limit(self):
        text, truncated = _enforce_word_limit("one two three", 5)
        assert text == "one two three"
        assert truncated is False

    def test_at_limit(self):
        text, truncated = _enforce_word_limit("one two three", 3)
        assert text == "one two three"
        assert truncated is False

    def test_over_limit(self):
        text, truncated = _enforce_word_limit("one two three four five", 3)
        assert text == "one two three"
        assert truncated is True

    def test_single_word_limit(self):
        text, truncated = _enforce_word_limit("hello world goodbye", 1)
        assert text == "hello"
        assert truncated is True

    def test_large_text_truncation(self):
        """500-word text truncated to 250 words."""
        words = [f"word{i}" for i in range(500)]
        text, truncated = _enforce_word_limit(" ".join(words), 250)
        assert truncated is True
        assert len(text.split()) == 250

    def test_zero_limit(self):
        """Zero word limit truncates to empty string."""
        text, truncated = _enforce_word_limit("hello world", 0)
        assert text == ""
        assert truncated is True


# ══════════════════════════════════════════════════════════════════════════
# POST-FILL VALIDATION TESTS
# ══════════════════════════════════════════════════════════════════════════


class TestValidateCompletedDocument:
    """Test the post-fill validation."""

    def test_identical_documents(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_formatted_template(f.name)
            try:
                warnings = _validate_completed_document(f.name, f.name)
                assert len(warnings) == 0
            finally:
                os.unlink(f.name)

    def test_filled_document_preserves_structure(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "Response", "word_limit": None}]
                )

                warnings = _validate_completed_document(tmp_in.name, output_path)
                assert len(warnings) == 0, f"Unexpected warnings: {warnings}"
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    def test_simple_fixture_validation(self, temp_output, simple_analysis):
        """Validate simple fixture after filling all fields."""
        mappings = [{
            "table_index": f["table_index"],
            "row_index": f["row_index"],
            "col_index": f["col_index"],
            "response_text": "Validation response",
            "word_limit": None,
        } for f in simple_analysis["fields"]]

        fill_template(SIMPLE_FIXTURE, temp_output, mappings)
        warnings = _validate_completed_document(SIMPLE_FIXTURE, temp_output)
        assert warnings == []

    def test_complex_fixture_validation(self, temp_output, complex_analysis):
        """Validate complex fixture after filling all fields."""
        mappings = [{
            "table_index": f["table_index"],
            "row_index": f["row_index"],
            "col_index": f["col_index"],
            "response_text": "Validation response",
            "word_limit": None,
        } for f in complex_analysis["fields"]]

        fill_template(COMPLEX_FIXTURE, temp_output, mappings)
        warnings = _validate_completed_document(COMPLEX_FIXTURE, temp_output)
        assert warnings == []


# ══════════════════════════════════════════════════════════════════════════
# TRACK CHANGES SAFE OPENING TESTS
# ══════════════════════════════════════════════════════════════════════════


class TestFillTemplateTrackChanges:
    """Tests that fill_template uses open_document_safe for TC handling."""

    @patch("fill_template.open_document_safe")
    def test_uses_open_document_safe(self, mock_open_safe):
        """fill_template should call open_document_safe instead of Document()."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            # Make open_document_safe return the real document (no TC)
            real_doc = Document(tmp_in.name)
            mock_open_safe.return_value = (real_doc, None)

            try:
                fill_template(
                    tmp_in.name, output_path,
                    [{"table_index": 0, "row_index": 1, "col_index": 1,
                      "response_text": "TC test", "word_limit": None}]
                )

                mock_open_safe.assert_called_once_with(tmp_in.name)
            finally:
                os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)

    @patch("fill_template.open_document_safe")
    def test_cleans_up_temp_file(self, mock_open_safe):
        """Temp file from TC resolution should be cleaned up after fill."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            _create_formatted_template(tmp_in.name)
            output_path = tmp_in.name.replace(".docx", "_out.docx")

            real_doc = Document(tmp_in.name)
            temp_path = "/tmp/fake_tc_resolved.docx"
            mock_open_safe.return_value = (real_doc, temp_path)

            try:
                with patch("os.unlink") as mock_unlink:
                    fill_template(
                        tmp_in.name, output_path,
                        [{"table_index": 0, "row_index": 1, "col_index": 1,
                          "response_text": "TC cleanup test", "word_limit": None}]
                    )

                    mock_unlink.assert_called_once_with(temp_path)
            finally:
                # Clean up manually since os.unlink was mocked
                if os.path.exists(tmp_in.name):
                    os.unlink(tmp_in.name)
                if os.path.exists(output_path):
                    os.unlink(output_path)


# ══════════════════════════════════════════════════════════════════════════
# PDF WRITER TESTS (ID-145 {145.15} — pypdf AcroForm value writes)
# ══════════════════════════════════════════════════════════════════════════


def _build_fillable_pdf(field_names: list[str]) -> bytes:
    """Programmatically build a minimal single-page fillable PDF with one
    ``/Tx`` Widget annotation per name, in the given order.

    Real fillable PDFs come from commonforms ({145.11}); this synthetic
    fixture avoids booting the real ML detector in a unit test while
    exercising the SAME page-major + Widget-annotation walk
    ``_pdf_field_names_by_sequence`` performs on commonforms' output.
    """
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        BooleanObject,
        DictionaryObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    field_refs = []
    for i, name in enumerate(field_names):
        field = DictionaryObject()
        field.update(
            {
                NameObject("/FT"): NameObject("/Tx"),
                NameObject("/T"): TextStringObject(name),
                NameObject("/Subtype"): NameObject("/Widget"),
                NameObject("/Rect"): ArrayObject(
                    [
                        NumberObject(x)
                        for x in (100, 700 - i * 40, 300, 720 - i * 40)
                    ]
                ),
                NameObject("/F"): NumberObject(4),
                NameObject("/Ff"): NumberObject(0),
            }
        )
        field_refs.append(writer._add_object(field))
    page[NameObject("/Annots")] = ArrayObject(field_refs)

    acroform = DictionaryObject()
    acroform[NameObject("/Fields")] = ArrayObject(field_refs)
    acroform[NameObject("/NeedAppearances")] = BooleanObject(True)
    writer._root_object[NameObject("/AcroForm")] = acroform

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


class TestPdfFieldNamesBySequence:
    """_pdf_field_names_by_sequence rebuilds sequence -> field_name in
    page-major + Widget-annotation order."""

    def test_maps_sequence_to_field_name_in_order(self):
        pdf_bytes = _build_fillable_pdf(["textbox_0_0", "textbox_0_1", "textbox_0_2"])
        result = _pdf_field_names_by_sequence(pdf_bytes)
        assert result == {0: "textbox_0_0", 1: "textbox_0_1", 2: "textbox_0_2"}


class TestFillPdfTemplate:
    """fill_pdf_template writes response text into AcroForm fields,
    addressing by reconstructed sequence -> field_name (form_instance_fields
    carries no field_name for PDF rows, {145.15})."""

    def test_fills_field_matched_by_sequence(self, tmp_path):
        pdf_bytes = _build_fillable_pdf(["textbox_0_0", "textbox_0_1"])
        input_path = str(tmp_path / "fillable.pdf")
        output_path = str(tmp_path / "out.pdf")
        with open(input_path, "wb") as f:
            f.write(pdf_bytes)

        result = fill_pdf_template(
            input_path,
            output_path,
            [
                {
                    "table_index": 0,
                    "row_index": 0,
                    "response_text": "Answer A",
                    "word_limit": None,
                }
            ],
        )

        assert result["fields_filled"] == 1
        assert result["fields_failed"] == 0
        reader = PdfReader(output_path)
        fields = reader.get_fields()
        assert fields["textbox_0_0"].get("/V") == "Answer A"

    def test_empty_response_skipped(self, tmp_path):
        pdf_bytes = _build_fillable_pdf(["textbox_0_0"])
        input_path = str(tmp_path / "fillable.pdf")
        output_path = str(tmp_path / "out.pdf")
        with open(input_path, "wb") as f:
            f.write(pdf_bytes)

        result = fill_pdf_template(
            input_path,
            output_path,
            [{"table_index": 0, "row_index": 0, "response_text": "", "word_limit": None}],
        )

        assert result["fields_skipped"] == 1
        assert result["fields_filled"] == 0

    def test_missing_sequence_fails(self, tmp_path):
        pdf_bytes = _build_fillable_pdf(["textbox_0_0"])
        input_path = str(tmp_path / "fillable.pdf")
        output_path = str(tmp_path / "out.pdf")
        with open(input_path, "wb") as f:
            f.write(pdf_bytes)

        result = fill_pdf_template(
            input_path,
            output_path,
            [
                {
                    "table_index": 0,
                    "row_index": 99,
                    "response_text": "Orphan mapping",
                    "word_limit": None,
                }
            ],
        )

        assert result["fields_failed"] == 1
        assert "No AcroForm widget" in result["errors"][0]["error"]

    def test_word_limit_truncation(self, tmp_path):
        pdf_bytes = _build_fillable_pdf(["textbox_0_0"])
        input_path = str(tmp_path / "fillable.pdf")
        output_path = str(tmp_path / "out.pdf")
        with open(input_path, "wb") as f:
            f.write(pdf_bytes)

        long_text = " ".join(["word"] * 20)
        result = fill_pdf_template(
            input_path,
            output_path,
            [
                {
                    "table_index": 0,
                    "row_index": 0,
                    "response_text": long_text,
                    "word_limit": 5,
                }
            ],
        )

        assert result["fields_filled"] == 1
        assert len(result["truncated"]) == 1
        assert result["truncated"][0]["limit"] == 5
        reader = PdfReader(output_path)
        written = reader.get_fields()["textbox_0_0"].get("/V")
        assert len(written.split()) == 5

    def test_refill_pass_preserves_prior_value_and_fills_gap(self, tmp_path):
        """A re-fill pass, using pass 1's OUTPUT as its base (BI-22
        re-entrancy — bid_worker.fill_template_job downloads the latest
        completion as the base on a subsequent pass), preserves the
        previously-filled field's value while filling a new gap. Proves
        the sequence -> field_name reconstruction stays stable across a
        pypdf write/read round-trip."""
        pdf_bytes = _build_fillable_pdf(["textbox_0_0", "textbox_0_1"])
        pass1_in = str(tmp_path / "p1_in.pdf")
        pass1_out = str(tmp_path / "p1_out.pdf")
        with open(pass1_in, "wb") as f:
            f.write(pdf_bytes)

        fill_pdf_template(
            pass1_in,
            pass1_out,
            [
                {
                    "table_index": 0,
                    "row_index": 0,
                    "response_text": "First answer",
                    "word_limit": None,
                }
            ],
        )

        pass2_out = str(tmp_path / "p2_out.pdf")
        result2 = fill_pdf_template(
            pass1_out,
            pass2_out,
            [
                {
                    "table_index": 0,
                    "row_index": 1,
                    "response_text": "Second answer",
                    "word_limit": None,
                }
            ],
        )

        assert result2["fields_filled"] == 1
        fields = PdfReader(pass2_out).get_fields()
        assert fields["textbox_0_0"].get("/V") == "First answer"
        assert fields["textbox_0_1"].get("/V") == "Second answer"


# ══════════════════════════════════════════════════════════════════════════
# XLSX WRITER TESTS (ID-145 {145.15} — openpyxl cell writes at recorded coords)
# ══════════════════════════════════════════════════════════════════════════

_WALK_SHEET_TARGET = "scripts.cocoindex_pipeline.form_extractors.xlsx._walk_sheet"


class TestXlsxSheetByTableIndex:
    """_xlsx_sheet_by_table_index reconstructs table_index -> sheet_name
    boundaries by re-running the deterministic per-sheet walk (xlsx.py's
    table_index is a workbook-wide counter, not a sheet index — no
    sheet_name column exists anywhere in the pipeline)."""

    def test_maps_table_index_across_multiple_sheets(self):
        wb = openpyxl.Workbook()
        wb.active.title = "Bidder 1"
        wb.create_sheet("Bidder 2")
        buf = io.BytesIO()
        wb.save(buf)
        raw_bytes = buf.getvalue()

        with patch(_WALK_SHEET_TARGET) as mock_walk:
            mock_walk.side_effect = [
                ([], 0, 2),  # "Bidder 1" consumes table_index 0, 1
                ([], 0, 1),  # "Bidder 2" consumes table_index 2
            ]
            result = _xlsx_sheet_by_table_index(raw_bytes)

        assert result == {0: "Bidder 1", 1: "Bidder 1", 2: "Bidder 2"}


class TestFillXlsxTemplate:
    """fill_xlsx_template writes response text into the correct sheet/cell,
    resolving sheet identity from table_index via the reconstructed
    boundary map (the multi-sheet addressing gap, {145.15})."""

    def _make_two_sheet_workbook(self, tmp_path) -> str:
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Bidder 1"
        ws1["A1"] = "Question"
        wb.create_sheet("Bidder 2")
        input_path = str(tmp_path / "in.xlsx")
        wb.save(input_path)
        return input_path

    def test_writes_to_correct_sheet_for_table_index(self, tmp_path):
        input_path = self._make_two_sheet_workbook(tmp_path)
        output_path = str(tmp_path / "out.xlsx")

        with patch(_WALK_SHEET_TARGET) as mock_walk:
            mock_walk.side_effect = [([], 0, 1), ([], 0, 1)]
            result = fill_xlsx_template(
                input_path,
                output_path,
                [
                    {
                        "table_index": 1,
                        "row_index": 1,
                        "col_index": 2,
                        "response_text": "Answer on Bidder 2",
                        "word_limit": None,
                    }
                ],
            )

        assert result["fields_filled"] == 1
        out_wb = openpyxl.load_workbook(output_path)
        assert out_wb["Bidder 2"].cell(row=1, column=2).value == "Answer on Bidder 2"
        assert out_wb["Bidder 1"].cell(row=1, column=2).value is None

    def test_missing_table_index_fails(self, tmp_path):
        input_path = self._make_two_sheet_workbook(tmp_path)
        output_path = str(tmp_path / "out.xlsx")

        with patch(_WALK_SHEET_TARGET) as mock_walk:
            mock_walk.side_effect = [([], 0, 1), ([], 0, 1)]
            result = fill_xlsx_template(
                input_path,
                output_path,
                [
                    {
                        "table_index": 99,
                        "row_index": 1,
                        "col_index": 2,
                        "response_text": "Orphan mapping",
                        "word_limit": None,
                    }
                ],
            )

        assert result["fields_failed"] == 1
        assert "No sheet found" in result["errors"][0]["error"]

    def test_empty_response_skipped(self, tmp_path):
        input_path = self._make_two_sheet_workbook(tmp_path)
        output_path = str(tmp_path / "out.xlsx")

        with patch(_WALK_SHEET_TARGET) as mock_walk:
            mock_walk.side_effect = [([], 0, 1), ([], 0, 1)]
            result = fill_xlsx_template(
                input_path,
                output_path,
                [
                    {
                        "table_index": 0,
                        "row_index": 1,
                        "col_index": 2,
                        "response_text": "   ",
                        "word_limit": None,
                    }
                ],
            )

        assert result["fields_skipped"] == 1
        assert result["fields_filled"] == 0

    def test_word_limit_truncation(self, tmp_path):
        input_path = self._make_two_sheet_workbook(tmp_path)
        output_path = str(tmp_path / "out.xlsx")
        long_text = " ".join(["word"] * 20)

        with patch(_WALK_SHEET_TARGET) as mock_walk:
            mock_walk.side_effect = [([], 0, 1), ([], 0, 1)]
            result = fill_xlsx_template(
                input_path,
                output_path,
                [
                    {
                        "table_index": 0,
                        "row_index": 1,
                        "col_index": 2,
                        "response_text": long_text,
                        "word_limit": 5,
                    }
                ],
            )

        assert result["fields_filled"] == 1
        assert len(result["truncated"]) == 1
        out_wb = openpyxl.load_workbook(output_path)
        written = out_wb["Bidder 1"].cell(row=1, column=2).value
        assert len(written.split()) == 5
