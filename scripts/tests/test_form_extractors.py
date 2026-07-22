"""Real-behaviour tests for the ID-52 form extractors (PRODUCT Inv-2),
generalised under DR-058 (ID-145.10).

NO mocks of python-docx / openpyxl internals, per
``docs/reference/testing/test-philosophy.md``.

PDF SCOPE NOTE (ID-145.10 Checker ruling, post-{145.11} landing): the
recovered id-52 ``pdf.py`` (pdfplumber AcroForm/table-walker) has been
REMOVED from this branch — {145.11} landed its own commonforms-based PDF
Plane-2 detector at the identical
``scripts/cocoindex_pipeline/form_extractors/pdf.py`` path (DR-057: real UK
procurement PDFs are FLAT, so detection via commonforms is mandatory, not
the AcroForm/table-walk this id-52 module attempted). PDF test coverage is
{145.11}'s own suite (``test_pdf_field_detection.py``); this file covers
DOCX + XLSX only.
"""

from __future__ import annotations

import asyncio
from pathlib import Path

import pytest
from docx import Document

from scripts.cocoindex_pipeline.form_extractors.docx import extract as docx_extract
from scripts.cocoindex_pipeline.form_extractors.shared import (
    ExtractedField,
    ExtractedForm,
    FormExtractionError,
    FormMetadata,
)
from scripts.cocoindex_pipeline.form_extractors.xlsx import extract as xlsx_extract

# ──────────────────────────────────────────────────────────────────────────
# Fixture path — committed symlinks to the canonical corpus files.
# ──────────────────────────────────────────────────────────────────────────

_FIXTURE_DIR = Path(__file__).parent / "fixtures" / "form-extraction"
_EFA_XLSX_PATH = _FIXTURE_DIR / "evaluation-matrix-itt-vol8.xlsx"


# ──────────────────────────────────────────────────────────────────────────
# Shared-model shape (TECH §2.2 strictness)
# ──────────────────────────────────────────────────────────────────────────


class TestSharedModels:
    """Smoke tests for the ``shared.py`` Pydantic shapes."""

    def test_extracted_field_forbids_extra_keys(self) -> None:
        """``extra="forbid"`` — surfaces drift at construction time."""
        from pydantic import ValidationError

        with pytest.raises(ValidationError):
            ExtractedField(
                field_type="empty_cell",
                fill_status="pending",
                sequence=0,
                unexpected_extra_key="boom",  # type: ignore[call-arg]
            )

    def test_extracted_field_strict_typing(self) -> None:
        """``strict=True`` — rejects coerced types per
        cocoindex-extraction-contract pattern."""
        from pydantic import ValidationError

        with pytest.raises(ValidationError):
            ExtractedField(
                field_type="empty_cell",
                fill_status="pending",
                sequence="0",  # type: ignore[arg-type]
            )

    def test_extracted_field_field_type_literal_set(self) -> None:
        """``field_type`` matches TECH §2.2 literal set exactly."""
        from pydantic import ValidationError

        with pytest.raises(ValidationError):
            ExtractedField(
                field_type="unknown_type",  # type: ignore[arg-type]
                fill_status="pending",
                sequence=0,
            )

    def test_extracted_field_fill_status_literal_set(self) -> None:
        """``fill_status`` matches TECH §2.2 literal set exactly."""
        from pydantic import ValidationError

        with pytest.raises(ValidationError):
            ExtractedField(
                field_type="empty_cell",
                fill_status="wrong",  # type: ignore[arg-type]
                sequence=0,
            )

    def test_extracted_field_defaults(self) -> None:
        """Reference URLs default to empty list; optionals default None."""
        field = ExtractedField(
            field_type="empty_cell",
            fill_status="pending",
            sequence=0,
        )
        assert field.reference_urls == []
        assert field.question_text is None
        assert field.placeholder_text is None
        assert field.row_index is None
        assert field.col_index is None
        assert field.table_index is None
        assert field.section_name is None
        assert field.word_limit is None
        assert field.is_mandatory is None

    def test_extracted_form_carries_metadata_and_fields(self) -> None:
        meta = FormMetadata(form_type="questionnaire", form_format="pdf")
        form = ExtractedForm(form_metadata=meta, fields=[])
        assert form.form_metadata.form_format == "pdf"
        assert form.fields == []

    def test_form_extraction_error_carries_reason_and_path(self) -> None:
        err = FormExtractionError(
            reason="corrupt_pdf", rel_path="foo/bar.pdf", details="boom"
        )
        assert err.reason == "corrupt_pdf"
        assert err.rel_path == "foo/bar.pdf"
        assert err.details == "boom"
        assert "corrupt_pdf" in str(err)
        assert "foo/bar.pdf" in str(err)
        assert "boom" in str(err)


# PDF-specific Inv-8/9/10/11/12/15/17 acceptance classes (TestPdfContainer
# ArtefactBypass, TestPdfMandatoryFlag, TestPdfWordLimit, TestPdfCoordinates,
# TestPdfSectionAndSequence, TestPdfPlaceholderVsAuthored,
# TestPdfFailureSurfacing) were REMOVED here (ID-145.10 Checker ruling,
# post-{145.11} landing) — their subject (the recovered id-52 pdfplumber
# reader) left this branch; PDF Plane-2 coverage is {145.11}'s own
# ``test_pdf_field_detection.py`` against the commonforms-based detector
# that now owns ``scripts/cocoindex_pipeline/form_extractors/pdf.py``.


# ══════════════════════════════════════════════════════════════════════════
# XLSX reader — {52.10} acceptance tests against the EFA fixture.
#
# Real-behaviour tests (per ``docs/reference/testing/test-philosophy.md``): the
# extractor reads the actual EFA fixture, NO mocks of openpyxl internals.
# The fixture lives under ``scripts/tests/fixtures/form-extraction`` as a
# symlink to ``docs/testing/test-data/templates/itt-services-efa/...``.
# ══════════════════════════════════════════════════════════════════════════


@pytest.fixture(scope="module")
def efa_xlsx_bytes() -> bytes:
    """Raw bytes for the EFA evaluation-matrix XLSX fixture."""
    assert _EFA_XLSX_PATH.exists(), (
        f"corpus fixture missing — {_EFA_XLSX_PATH} should symlink to "
        f"docs/testing/test-data/templates/itt-services-efa/"
        f"evaluation-matrix-itt-vol8.xlsx"
    )
    return _EFA_XLSX_PATH.read_bytes()


@pytest.fixture(scope="module")
def efa_form(efa_xlsx_bytes: bytes) -> ExtractedForm:
    """Module-scoped EFA extraction — amortise openpyxl walk."""
    return asyncio.run(
        xlsx_extract(efa_xlsx_bytes, "evaluation-matrix-itt-vol8.xlsx")
    )


# ──────────────────────────────────────────────────────────────────────────
# Inv-2 / TECH §2.2 — public shape matches PDF reader
# ──────────────────────────────────────────────────────────────────────────


class TestXlsxPublicShape:
    """The XLSX extractor exposes the same ``extract(raw_bytes, filename)
    -> ExtractedForm`` async signature as the PDF reader (TECH §2.2)."""

    def test_efa_extracts_to_extracted_form(self, efa_form: ExtractedForm) -> None:
        assert isinstance(efa_form, ExtractedForm)
        assert isinstance(efa_form.form_metadata, FormMetadata)
        assert efa_form.form_metadata.form_format == "xlsx"
        assert efa_form.fields, "no fields extracted from EFA"
        for field in efa_form.fields:
            assert isinstance(field, ExtractedField)


# ──────────────────────────────────────────────────────────────────────────
# Inv-13 — per-form dedup: Bidder 1 ≡ Bidder 2 → N fields, NOT 2N
# ──────────────────────────────────────────────────────────────────────────


class TestXlsxEfaDedup:
    """PRODUCT Inv-13 — the EFA workbook repeats the same scoring matrix
    across the ``Bidder 1`` and ``Bidder 2`` sheets; the extractor records
    each distinct question once, not once per copy. Dedup is keyed on
    ``(section_name, normalise(question_text))`` per TECH §2.3."""

    def test_efa_total_field_count_deduped_to_19(
        self, efa_form: ExtractedForm
    ) -> None:
        """Inv-13 — the EFA workbook repeats its scoring matrix across the
        ``Bidder 1`` and ``Bidder 2`` sheets. With per-form dedup keyed on
        ``(section_name, normalise(question_text))`` the extractor records
        each distinct question once: 19 deduped fields (Part2=2, Part3=1,
        Part4=5, Part5=3, Part6=1, Part7=7), NOT 38 (19 per bidder copy)."""
        assert len(efa_form.fields) == 19, (
            f"Inv-13 EFA dedup regression — expected exactly 19 deduped "
            f"fields (Bidder 1 ≡ Bidder 2 collapsed), got "
            f"{len(efa_form.fields)}"
        )

    def test_efa_bidder_questions_deduped_to_n_not_2n(
        self, efa_form: ExtractedForm
    ) -> None:
        """Every ``2.1`` / ``3.1`` / ``4.1..4.5`` / ``5.1..5.3`` /
        ``6.1`` / ``7.1..7.7`` row appears EXACTLY once across the
        whole form, not twice (once per bidder sheet)."""
        authored = [
            f
            for f in efa_form.fields
            if f.question_text and "Bidders" in (f.question_text or "")
        ] + [
            f
            for f in efa_form.fields
            if f.question_text and any(
                tag in f.question_text
                for tag in (
                    "Overall assessment",
                    "Detailed programme",
                    "Approach to construction",
                    "Management procedure",
                    "Panel Members",
                    "Total project cost",
                    "Pricing schedules",
                    "Life cycle",
                )
            )
        ]
        # Group by normalised text — every group should be size 1 (deduped).
        groups: dict[str, int] = {}
        for f in authored:
            key = (f.question_text or "").strip().lower()
            groups[key] = groups.get(key, 0) + 1
        duplicates = {k: c for k, c in groups.items() if c > 1}
        assert not duplicates, (
            f"EFA dedup regression — {len(duplicates)} question(s) appear "
            f"more than once across Bidder 1 / Bidder 2: "
            f"{list(duplicates.items())[:5]}"
        )

    def test_efa_known_question_text_present_exactly_once(
        self, efa_form: ExtractedForm
    ) -> None:
        """The 2.1 question (\"Bidders should outline their overall
        approach…\") exists on BOTH Bidder 1 (rendered with a typo
        \"nsample\") and Bidder 2 (rendered \"sample\") in the raw
        workbook. Per Inv-13 dedup keyed on the NORMALISED text these
        are distinct, so the extractor preserves both as distinct
        candidates. The acceptance bar is that within each rendering,
        the question appears once — not duplicated by sheet."""
        approach = [
            f
            for f in efa_form.fields
            if f.question_text and "overall approach" in (f.question_text.lower())
        ]
        # Each surface variant should appear once.
        normalised: dict[str, int] = {}
        for f in approach:
            key = (f.question_text or "").strip().lower()
            normalised[key] = normalised.get(key, 0) + 1
        for variant, count in normalised.items():
            assert count == 1, (
                f"variant {variant!r} appeared {count} times — duplicate "
                f"sheet copy not deduped"
            )


# ──────────────────────────────────────────────────────────────────────────
# Inv-8 — coordinates populated for scoring-matrix rows
# ──────────────────────────────────────────────────────────────────────────


class TestXlsxEfaCoordinates:
    """PRODUCT Inv-8 — scoring-matrix rows expose row/col/table indices."""

    def test_efa_scoring_matrix_rows_have_coordinates(
        self, efa_form: ExtractedForm
    ) -> None:
        scoring_fields = [
            f
            for f in efa_form.fields
            if f.section_name and "OVERALL APPROACH" in (f.section_name or "").upper()
        ]
        assert scoring_fields, "no Part 2 OVERALL APPROACH question found"
        for f in scoring_fields:
            assert isinstance(f.row_index, int), (
                f"row_index missing on {f.question_text!r}"
            )
            assert isinstance(f.col_index, int), (
                f"col_index missing on {f.question_text!r}"
            )
            assert isinstance(f.table_index, int), (
                f"table_index missing on {f.question_text!r}"
            )


# ──────────────────────────────────────────────────────────────────────────
# Inv-12 — section names: Part 2 — OVERALL APPROACH etc.
# ──────────────────────────────────────────────────────────────────────────


class TestXlsxEfaSections:
    """PRODUCT Inv-12 — section hierarchy preserved from the
    Scoring Matrix's section-header rows (Part N — TITLE)."""

    def test_efa_part_2_section_name_recorded(
        self, efa_form: ExtractedForm
    ) -> None:
        part2 = [
            f
            for f in efa_form.fields
            if f.section_name and "OVERALL APPROACH" in (f.section_name or "").upper()
        ]
        assert part2, (
            "no Part 2 OVERALL APPROACH-sectioned field recorded — "
            "Inv-12 regression"
        )

    def test_efa_multiple_part_sections_observed(
        self, efa_form: ExtractedForm
    ) -> None:
        """The EFA scoring matrix has Parts 2, 3, 4, 5, 6, 7 — at least
        4 of those should appear as distinct section names."""
        section_names = {
            f.section_name
            for f in efa_form.fields
            if f.section_name and "PART" in (f.section_name or "").upper()
        }
        assert len(section_names) >= 4, (
            f"only {len(section_names)} distinct Part sections seen — "
            f"Inv-12 regression"
        )


# ──────────────────────────────────────────────────────────────────────────
# Inv-12 — sequence is reading-order
# ──────────────────────────────────────────────────────────────────────────


class TestXlsxSequence:
    """PRODUCT Inv-12 — reading-order ``sequence`` strictly increasing."""

    def test_efa_sequence_strictly_increasing(
        self, efa_form: ExtractedForm
    ) -> None:
        seqs = [f.sequence for f in efa_form.fields]
        assert seqs == sorted(seqs), (
            "EFA sequence not in reading order — Inv-12 regression"
        )
        assert len(set(seqs)) == len(seqs), (
            "EFA duplicate sequence values — Inv-12 regression"
        )


# ──────────────────────────────────────────────────────────────────────────
# Inv-17 — corrupt / empty XLSX raises FormExtractionError
# ──────────────────────────────────────────────────────────────────────────


class TestXlsxFailureSurfacing:
    """PRODUCT Inv-17 — corrupt input raises ``FormExtractionError``;
    never silently returns an empty ``ExtractedForm``."""

    def test_corrupt_bytes_raises_form_extraction_error(self) -> None:
        with pytest.raises(FormExtractionError) as excinfo:
            asyncio.run(xlsx_extract(b"not a real xlsx", "corrupt.xlsx"))
        assert excinfo.value.rel_path == "corrupt.xlsx"
        assert excinfo.value.reason

    def test_empty_bytes_raises_form_extraction_error(self) -> None:
        with pytest.raises(FormExtractionError) as excinfo:
            asyncio.run(xlsx_extract(b"", "empty.xlsx"))
        assert excinfo.value.rel_path == "empty.xlsx"
        assert excinfo.value.reason


# ──────────────────────────────────────────────────────────────────────────
# TECH §2.3 — dedup only collapses identical text WITHIN a single form
# ──────────────────────────────────────────────────────────────────────────


class TestXlsxDedupScope:
    """TECH §2.3 — dedup is per-form (workspace-scoped). Two distinct
    forms in the same workspace may legitimately share question text;
    catalogue-level reuse is Path-C's concern.

    This test verifies the in-form dedup does not flatten genuinely
    different questions: the EFA scoring matrix contains both a 4.1
    question (\"Overall assessment of Design solution\") and a 4.2
    question (\"Overall assessment of environmental strategy/solution\")
    — these share a 3-word prefix but are different questions and must
    BOTH survive dedup as separate fields.
    """

    def test_efa_similar_questions_not_collapsed(
        self, efa_form: ExtractedForm
    ) -> None:
        design = [
            f
            for f in efa_form.fields
            if f.question_text
            and "Overall assessment of Design solution"
            in (f.question_text or "")
        ]
        env = [
            f
            for f in efa_form.fields
            if f.question_text
            and "Overall assessment of environmental"
            in (f.question_text or "")
        ]
        assert design and env, (
            "4.1 / 4.2 'Overall assessment of …' questions missing — "
            "TECH §2.3 regression (over-collapsed two distinct questions)"
        )


# ──────────────────────────────────────────────────────────────────────────
# DOCX reader — Charnwood ``ITT Services.docx`` (PLAN §{52.11}).
# 1908 paragraphs + 8 tables; the question source is BOTH the prose
# (placeholder spans inside authored paragraphs) AND the tables
# (authored Q/A + placeholder grids).
# ──────────────────────────────────────────────────────────────────────────

_CHARNWOOD_DOCX_PATH = _FIXTURE_DIR / "itt-services-charnwood.docx"


@pytest.fixture(scope="module")
def charnwood_docx_bytes() -> bytes:
    """Raw bytes for the Charnwood DOCX fixture (real corpus file)."""
    assert _CHARNWOOD_DOCX_PATH.exists(), (
        f"corpus fixture missing — {_CHARNWOOD_DOCX_PATH} should symlink to "
        f"docs/testing/test-data/templates/itt-services-charnwood/"
        f"ITT Services.docx"
    )
    return _CHARNWOOD_DOCX_PATH.read_bytes()


@pytest.fixture(scope="module")
def charnwood_form(charnwood_docx_bytes: bytes) -> ExtractedForm:
    """Module-scoped extraction so the pandoc + python-docx walk runs
    once across the assertion bundle below (the Charnwood file has
    unaccepted tracked changes — pandoc resolves them on each open)."""
    return asyncio.run(
        docx_extract(charnwood_docx_bytes, "itt-services-charnwood.docx")
    )


class TestDocxFixtureBaseline:
    """Sanity-check the corpus shape the spec was written against."""

    def test_charnwood_raw_para_and_table_counts(
        self, charnwood_docx_bytes: bytes
    ) -> None:
        """PLAN §{52.11} acceptance line — 1908 paragraphs + 8 tables.

        Verifies the **raw** counts (no pandoc clean-up): the
        extractor operates against these bytes, and the spec count
        is the single point of reference for "the question source"
        the extractor must cover.
        """
        import io

        raw_doc = Document(io.BytesIO(charnwood_docx_bytes))
        assert len(raw_doc.paragraphs) == 1908, (
            f"raw paragraph count drifted from spec baseline (PLAN §{{52.11}}): "
            f"{len(raw_doc.paragraphs)} vs 1908"
        )
        assert len(raw_doc.tables) == 8, (
            f"raw table count drifted from spec baseline (PLAN §{{52.11}}): "
            f"{len(raw_doc.tables)} vs 8"
        )


class TestCharnwoodPerTable:
    """ID-52.20 investigation, SUPERSEDED by DR-058 (ID-145.10) — pin the
    field count of EVERY Charnwood table so no table's zero is silent.

    The original ID-52.20 ruling below pinned 7 of 8 Charnwood tables to a
    documented "legit-zero" because the id-52 readers modelled exactly two
    archetypes (an authored Q/A header, and a first-column placeholder
    grid) and nothing else. FORM-EXTRACTION-SPIKE.md §1.2 measured that
    directly as a DESIGN LIMIT, not a bounded bug — 7/8 real fillable
    Charnwood tables (contact block, timetable, checklist, pricing,
    declaration, signature) were silently dropped. TECH.md §3.2 / spike §5
    mandate generalising past the two archetypes: DR-058's shape-3 fallback
    (a generic labelled-cell -> empty/placeholder-cell detector, plus
    standalone-bracket-label detection) now recovers those tables. This
    class re-rules each table against the MEASURED post-generalisation
    output — every table is asserted to an EXACT count, and every
    remaining zero carries a documented reason. No silent
    under-*or*-over-extraction.

    Rulings (raw shape → reason):

    * ``table[0]`` 3×2 ``[Name]``/``[Tel]`` — **4 fields (shape 3, rule
      1 — standalone bracket-label cells).** ``[Name]``, ``[Tel]``,
      ``[Job title]``, ``[Departmental Email]`` each emit as their own
      ``field_type='placeholder'`` field (row2's ``[Your Department
      Name]`` is a horizontal merge — the merge-continuation guard emits
      it once, not twice). The ID-52.20 "legit-zero" ruling for this
      table is exactly the overfit DR-058 supersedes: a bare bracketed
      label IS a fill-in-the-blank slot, just not the narrower
      ``[Insert …]`` instruction shape the archetype patterns recognise.
    * ``table[1]`` 12×3 ``Stage``/``Date(s)`` — **10 fields (shape 3,
      rule 2 — adjacent label + placeholder-cell pair).** The header row
      itself (``Stage``/``Date(s) and time(s)``) correctly emits nothing
      (neither cell is empty/placeholder-shaped); each data row's stage
      label pairs with its ``[Insert date]``/``[Insert date and time]``
      placeholder cell.
    * ``table[2]`` 12×2 ``SCHEDULE HEADING``/``COMPLETED?`` — **11 fields
      (shape 3, rule 2).** Each schedule label pairs with its ``□``
      checkbox-glyph answer cell (the generic answer-marker set admits
      checkbox glyphs alongside true emptiness and the archetype
      placeholder patterns).
    * ``table[3]`` 7×2 ``Insert question title``/``Insert %`` — **7
      fields, UNCHANGED (shape 2, the archetype placeholder-grid
      fast-path — DR-058 keeps the two archetypes as high-precision
      fast-paths layered OVER the generic detector).**
    * ``table[4]`` 5×2 ``0-3``/``Completely unsatisfactory …`` —
      **legit-zero, UNCHANGED.** Evaluator scoring rubric: BOTH cells on
      every row carry authored descriptor prose (a score band + its
      description), so shape 3's rule 2 correctly does not fire (it
      requires exactly one side of a pair to be empty/placeholder-shaped
      — two authored-prose cells is a reference table, not a fillable
      slot). The spike's own verdict calls this table "arguably correct"
      to leave at zero (FORM-EXTRACTION-SPIKE.md §1.2).
    * ``table[5]`` 8×2 ``Service component description``/``Costs (£)`` —
      **1 field (shape 3, rule 2 — the ``Total Costs (£) *``/``£`` row
      only).** The header row and the 6 blank line-item rows (both
      cells empty on every one) correctly emit nothing — rule 2 requires
      a labelled cell on one side, and a row with BOTH cells blank has no
      label to anchor a field on. Capturing the free-form per-line-item
      pricing grid would need a header-driven "column labels + blank body
      rows" pattern distinct from cell-adjacency pairing; left as a
      known, disclosed gap (not attempted — DR-058 scope is the
      cell-adjacency generalisation TECH.md §3.2 names, not a bespoke
      third archetype).
    * ``table[6]`` 6×2 ``I DECLARE THAT …``/``Name:`` … — **5 fields
      (shape 3, rule 2 — bidirectional).** The declaration statement row
      (both cells carry the same merged prose) correctly emits nothing;
      the 5 labelled slots (``Name:``, ``Position (Job Title):``,
      ``Date:``, ``Telephone number:``, ``Signature:``) each pair with
      their EMPTY left-hand cell — rule 2 checks both neighbour
      directions, since the blank slot sits to the label's LEFT here
      (unlike table[1]/[2]/[7] where it sits to the right).
    * ``table[7]`` 15×2 ``SIGNED for and on behalf of the Council``/`` ``
      — **10 fields (shape 3, rule 2).** Each signature-block label
      (``SIGNED for and on behalf of the Council``, ``Print Name and
      Address``, …) pairs with its empty right-hand cell; blank spacer
      rows (both cells empty) correctly emit nothing.

    FINAL Charnwood counts (DR-058 re-pin, supersedes {52.19}):
    total = 99, paragraph-sourced (``table_index is None``) = 51
    (UNCHANGED from {52.19} — the paragraph-level placeholder walk is
    untouched by this generalisation), table-sourced = 48 across 7 of
    the document's 8 tables (up from 7 fields / 1 table).
    """

    # table_index → (expected field count, ruling/reason). Zero entries
    # carry the documented reason as the assertion message so a future
    # drift cannot silently re-introduce an unexplained zero.
    _EXPECTED: dict[int, tuple[int, str]] = {
        0: (4, "shape 3 rule 1: standalone bracket-label cells ([Name]/[Tel]/[Job title]/[Departmental Email])"),
        1: (10, "shape 3 rule 2: timetable label + [Insert date] placeholder pairs"),
        2: (11, "shape 3 rule 2: schedule label + checkbox-glyph answer pairs"),
        3: (7, "placeholder grid (Inv-9, shape 2 archetype, unchanged): 7 'Insert question title' rows"),
        4: (0, "legit-zero, unchanged: evaluator scoring rubric — both cells are authored prose, no empty/placeholder side"),
        5: (1, "shape 3 rule 2: only the 'Total Costs (£) *' / '£' row — blank per-line-item rows have no labelled cell to anchor on (disclosed gap)"),
        6: (5, "shape 3 rule 2 (bidirectional): Name:/Position:/Date:/Telephone number:/Signature: pair with their EMPTY left-hand cell"),
        7: (10, "shape 3 rule 2: signature-block labels pair with their empty right-hand cell"),
    }

    def test_charnwood_per_table_field_count(
        self, charnwood_form: ExtractedForm, charnwood_docx_bytes: bytes
    ) -> None:
        """Every one of the 8 Charnwood tables is asserted to an EXACT
        expected count — non-zero with a captured shape, or zero with a
        documented legit-zero reason. No silent under-extraction."""
        import io

        counts: dict[int, int] = {ti: 0 for ti in self._EXPECTED}
        for field in charnwood_form.fields:
            if field.table_index is not None:
                counts[field.table_index] = counts.get(field.table_index, 0) + 1

        # Every table the document carries must be ruled (no table is
        # absent from the expectation map — guards against a new table
        # appearing and silently surfacing/dropping fields).
        raw_doc = Document(io.BytesIO(charnwood_docx_bytes))
        assert len(raw_doc.tables) == len(self._EXPECTED), (
            f"Charnwood table count drifted ({len(raw_doc.tables)} tables) — "
            f"the per-table ruling map covers {len(self._EXPECTED)}; re-rule "
            f"the new/removed table before re-pinning {{52.19}} counts"
        )

        for table_index, (expected, reason) in self._EXPECTED.items():
            actual = counts.get(table_index, 0)
            assert actual == expected, (
                f"table[{table_index}] field count {actual} != expected "
                f"{expected} — ruling: {reason}"
            )

    def test_charnwood_total_paragraph_table_counts(
        self, charnwood_form: ExtractedForm
    ) -> None:
        """DR-058 re-pin (supersedes {52.19}) — headline totals against the
        MEASURED post-generalisation output: 99 total, 51 paragraph-sourced
        (unchanged — the paragraph walk is untouched by DR-058), 48
        table-sourced across 7 of the document's 8 tables (up from 7 fields
        confined to table[3] alone). See TestCharnwoodPerTable for the
        per-table ruling this total is built from."""
        fields = charnwood_form.fields
        paragraph_sourced = [f for f in fields if f.table_index is None]
        table_sourced = [f for f in fields if f.table_index is not None]
        assert len(fields) == 99, (
            f"Charnwood TOTAL field count drifted: {len(fields)} != 99 "
            f"(DR-058 re-pin, feeds ID-145.10 regression gate)"
        )
        assert len(paragraph_sourced) == 51, (
            f"Charnwood paragraph-sourced count drifted: "
            f"{len(paragraph_sourced)} != 51 — the paragraph-level walk is "
            f"untouched by the DR-058 table generalisation"
        )
        assert len(table_sourced) == 48, (
            f"Charnwood table-sourced count drifted: "
            f"{len(table_sourced)} != 48"
        )
        tables_with_fields = {f.table_index for f in table_sourced}
        assert tables_with_fields == {0, 1, 2, 3, 5, 6, 7}, (
            f"Charnwood tables-with-fields drifted: {sorted(tables_with_fields)} "
            f"!= [0,1,2,3,5,6,7] — table[4] (the scoring rubric) is the one "
            f"documented legit-zero (TestCharnwoodPerTable); DR-058's "
            f"regression gate is 'Charnwood >= its 8 real fillable tables' — "
            f"7 of 8 clear that bar, up from 1 of 8 pre-generalisation"
        )


class TestDocxParasAndTablesBoth:
    """PRODUCT Inv-8 + Inv-12 — the extractor walks BOTH paragraphs AND
    tables, recording fields from each in reading-order ``sequence``."""

    def test_fields_originate_from_both_paragraphs_and_tables(
        self, charnwood_form: ExtractedForm
    ) -> None:
        """At least one field must come from a paragraph (table_index
        is None) AND at least one must come from a table (table_index
        is an integer). If only one side fires the extractor is
        dropping content the spec demands."""
        paragraph_sourced = [
            f for f in charnwood_form.fields if f.table_index is None
        ]
        table_sourced = [
            f for f in charnwood_form.fields if f.table_index is not None
        ]
        assert paragraph_sourced, (
            "no paragraph-sourced fields — the extractor missed the "
            "1908-paragraph prose-side question source (PLAN §{52.11} "
            "Inv-8 paras+tables)"
        )
        assert table_sourced, (
            "no table-sourced fields — the extractor missed the 8-table "
            "structural question source (PLAN §{52.11} Inv-8 paras+tables)"
        )

    def test_sequence_is_strictly_increasing(
        self, charnwood_form: ExtractedForm
    ) -> None:
        """Inv-12 — ``sequence`` must be the reading-order position
        across the WHOLE form (paragraphs and tables interleaved)."""
        seqs = [f.sequence for f in charnwood_form.fields]
        assert seqs == sorted(seqs), "sequence not in reading order"
        assert len(set(seqs)) == len(seqs), "duplicate sequence values"


class TestDocxPlaceholderGrid:
    """PRODUCT Inv-9 — ``Insert question title`` grid rows surface as
    placeholders, not authored questions.

    Charnwood ``ITT Services.docx`` table 3 is a 7-row 2-col grid:

    | Insert question title | Insert % |
    | Insert questions title | Insert % |
    | ... |

    Every first-column cell is placeholder text; the spec demands
    ``field_type='placeholder'``, ``placeholder_text`` populated,
    ``question_text=None`` per Inv-9.
    """

    def test_placeholder_grid_rows_have_no_authored_question(
        self, charnwood_form: ExtractedForm
    ) -> None:
        placeholder_grid_fields = [
            f
            for f in charnwood_form.fields
            if f.field_type == "placeholder"
            and f.placeholder_text
            and "Insert question" in f.placeholder_text
        ]
        assert placeholder_grid_fields, (
            "no 'Insert question title' grid rows extracted — Inv-9 "
            "placeholder-grid regression (Charnwood table 3)"
        )
        for f in placeholder_grid_fields:
            assert f.question_text is None, (
                f"placeholder-grid row carried question_text={f.question_text!r} — "
                f"Inv-9 violation (placeholder grid must not synthesise "
                f"an authored question)"
            )
            assert f.placeholder_text is not None
            assert f.field_type == "placeholder"
            assert f.table_index is not None, (
                "placeholder-grid row must carry table_index — Inv-8"
            )


class TestDocxAuthoredQuestionsPreserved:
    """PRODUCT Inv-9 — authored questions with empty answer cells must
    be preserved (not dropped because the answer slot is blank)."""

    def test_authored_paragraph_question_text_populated(
        self, charnwood_form: ExtractedForm
    ) -> None:
        """Paragraph-sourced placeholder fields must record the full
        paragraph as ``question_text`` (so the filler has context) and
        the placeholder span as ``placeholder_text``."""
        paragraph_sourced = [
            f
            for f in charnwood_form.fields
            if f.table_index is None
            and f.question_text
            and f.placeholder_text
        ]
        assert paragraph_sourced, (
            "no paragraph fields carrying both authored prose + "
            "placeholder span — Inv-9 paragraph-side regression"
        )
        # E-Mail: [Insert departmental email address] is in the
        # document header — confirm the placeholder span survives.
        email_field = [
            f
            for f in paragraph_sourced
            if f.placeholder_text
            and "departmental email" in f.placeholder_text.lower()
        ]
        assert email_field, (
            "expected '[Insert departmental email address]' placeholder "
            "from the document header — Inv-9 paragraph-side regression"
        )


class TestDocxWordLimit:
    """PRODUCT Inv-11 — word limit extracted via the reused
    ``_extract_word_limit`` helper when the form expresses one."""

    def test_extracted_word_limit_uses_reused_helper(self) -> None:
        """Direct unit check of the reused helper — ensures the import
        path delivers the same regex as the prior art."""
        from scripts.cocoindex_pipeline.form_extractors.docx import (
            _extract_word_limit as docx_extract_word_limit,
        )

        assert docx_extract_word_limit("Max 500 words") == 500
        assert docx_extract_word_limit("(no more than 250 words)") == 250
        assert docx_extract_word_limit("no limit stated") is None


class TestDocxSectionFromHeading:
    """PRODUCT Inv-12 — section name inherits from the most recent
    Heading-styled paragraph above the field.

    The reused ``_extract_section_headings`` helper maps tables to
    headings; our wrapper inlines that walk so paragraphs share the
    same streaming-section state.
    """

    def test_some_fields_carry_section_name(
        self, charnwood_form: ExtractedForm
    ) -> None:
        """At least some fields must carry a non-None ``section_name``
        — the Charnwood document uses Heading-styled paragraphs (e.g.
        ``SECTION A Company Details``)."""
        with_section = [
            f for f in charnwood_form.fields if f.section_name
        ]
        assert with_section, (
            "no fields carry section_name — Inv-12 regression (heading "
            "tracking broken)"
        )


class TestDocxMetadata:
    """``FormMetadata`` shape matches the PDF + XLSX readers."""

    def test_form_format_is_docx(self, charnwood_form: ExtractedForm) -> None:
        assert charnwood_form.form_metadata.form_format == "docx"

    def test_form_title_falls_back_to_filename(
        self, charnwood_form: ExtractedForm
    ) -> None:
        """When no title-styled paragraph is detected, the wrapper
        falls back to the supplied filename (mirrors the PDF reader's
        title fallback)."""
        assert charnwood_form.form_metadata.form_title == (
            "itt-services-charnwood.docx"
        )


class TestDocxFailureSurfacing:
    """PRODUCT Inv-17 — corrupt or empty input raises
    ``FormExtractionError``; never silently returns an empty
    ``ExtractedForm``."""

    def test_corrupt_bytes_raises_form_extraction_error(self) -> None:
        with pytest.raises(FormExtractionError) as excinfo:
            asyncio.run(docx_extract(b"not a real docx", "corrupt.docx"))
        assert excinfo.value.rel_path == "corrupt.docx"
        assert excinfo.value.reason  # non-empty machine-readable token

    def test_empty_bytes_raises_form_extraction_error(self) -> None:
        with pytest.raises(FormExtractionError) as excinfo:
            asyncio.run(docx_extract(b"", "empty.docx"))
        assert excinfo.value.rel_path == "empty.docx"
        assert excinfo.value.reason == "empty_docx"


# TestCorruptPdfBatchFixture (Inv-17: the committed `corrupt.pdf` batch
# fixture raises via the PDF reader) was REMOVED here for the same reason
# as the PDF acceptance classes above — its subject was the id-52
# pdfplumber reader's error path, which no longer exists on this branch.
# {145.11}'s commonforms-based detector owns corrupt-PDF error surfacing
# now; the `corrupt.pdf` fixture file itself is untouched (still committed,
# still available to {145.11}/{145.13}'s own test suites).


# ──────────────────────────────────────────────────────────────────────────
# DR-058 (ID-145.10) — generalisation regression gate. The two archetypes
# (authored Q/A header, placeholder grid for DOCX; CSP/EFA header signatures
# for XLSX) reproduce the id-52 ACCEPTANCE.md numbers exactly on in-corpus
# fixtures but measured ZERO on unseen real forms
# (FORM-EXTRACTION-SPIKE.md §1.1). This section pins the corpus regression
# gate the recovery+generalisation subtask brief names explicitly: Charnwood
# >= 8 real fillable tables (TestCharnwoodPerTable, above), annex_2 > 0,
# annex_3 rate-card rows detected, EFA stays exactly 19 (no regression from
# the new XLSX fallback), CSP stays exactly 45 (the archetype fast-path,
# unchanged by generalisation) — plus two isolated synthetic proofs that the
# generic rule itself (not just the real corpus) fires beyond the two
# archetypes, and one proof that it does NOT fire on non-form metadata
# (the measured false-positive risk the header+repetition gating in
# xlsx.py exists to close — EFA's own "Title Sheet"/"Summary" sheets).
# ──────────────────────────────────────────────────────────────────────────


class TestDR058UnseenRealFormsRegressionGate:
    """Real-corpus regression gate, measured against the owner-provided
    forms the id-52 archetypes never saw."""

    _ANNEX_2_PATH = _FIXTURE_DIR / "annex_2_supplier_response.docx"
    _ANNEX_3_PATH = _FIXTURE_DIR / "annex_3_pricing_approach.xlsx"
    _CSP_PATH = _FIXTURE_DIR / "Cloud Security Principles Checklist V5_3.xlsx"

    def test_annex_2_docx_yields_nonzero_fields(self) -> None:
        """British Council annex_2 supplier-response DOCX — measured ZERO
        pre-generalisation (0 fields; FORM-EXTRACTION-SPIKE.md §1.1's
        header row is a merged section banner, not the id-52 archetypes'
        assumed ``rows[0]``). The requirement prose lives one cell per row,
        followed by trailing blank paragraphs after a "Supplier Response:"
        label — the shape-3 rule-3 cell-internal-trailing-blank detector."""
        assert self._ANNEX_2_PATH.exists(), (
            f"fixture missing — {self._ANNEX_2_PATH} should symlink to "
            f"docs/testing/test-data/templates/rfp-british-council/"
            f"annex_2_supplier_response.docx"
        )
        raw = self._ANNEX_2_PATH.read_bytes()
        form = asyncio.run(docx_extract(raw, "annex_2_supplier_response.docx"))
        assert len(form.fields) > 0, (
            "annex_2 must yield >0 fields post-generalisation (measured 0 "
            "pre-generalisation — FORM-EXTRACTION-SPIKE.md §1.1)"
        )
        # At least one field must carry the real requirement prose (not
        # just header-row noise) — proves rule 3 actually fired.
        assert any(
            f.question_text and "avoided" in f.question_text.lower()
            for f in form.fields
        ), "expected the Social Value requirement prose among the extracted fields"

    def test_annex_3_xlsx_rate_card_rows_detected(self) -> None:
        """British Council annex_3 rate-card XLSX — measured ZERO
        pre-generalisation (FORM-EXTRACTION-SPIKE.md §1.1). The "Rate Card
        & Resources" sheet matches neither the EFA nor CSP header
        signature; each role row pairs a labelled Role cell with its
        unfilled Day Rate cell (a bare numeric zero default)."""
        assert self._ANNEX_3_PATH.exists(), (
            f"fixture missing — {self._ANNEX_3_PATH} should symlink to "
            f"docs/testing/test-data/templates/rfp-british-council/"
            f"annex_3_pricing_approach.xlsx"
        )
        raw = self._ANNEX_3_PATH.read_bytes()
        form = asyncio.run(xlsx_extract(raw, "annex_3_pricing_approach.xlsx"))
        assert len(form.fields) > 0, (
            "annex_3 must yield >0 fields post-generalisation (measured 0 "
            "pre-generalisation — FORM-EXTRACTION-SPIKE.md §1.1)"
        )
        # At least one field must be a rate-card role row (not just the
        # VAT/Overall Price summary lines).
        assert any(
            f.question_text and f.question_text.lower().startswith("e.g.")
            for f in form.fields
        ), "expected at least one 'e.g. <role>' rate-card row among the extracted fields"

    def test_csp_xlsx_matches_acceptance_baseline_unregressed(self) -> None:
        """CSP Cloud Security Principles checklist — the CSP archetype
        fast-path (unchanged by DR-058) must still hit its measured 45
        (FORM-EXTRACTION-SPIKE.md §1.1 / the id-52 ACCEPTANCE.md baseline);
        the generic XLSX fallback must NOT also fire on this sheet (it is
        archetype-claimed, so shape 3 never runs on it — a drift here would
        mean the archetype dispatch order broke)."""
        assert self._CSP_PATH.exists(), (
            f"fixture missing — {self._CSP_PATH} should symlink to "
            f"docs/testing/test-data/templates/csp-cloud-security-principles/"
            f"Cloud Security Principles Checklist V5_3.xlsx"
        )
        raw = self._CSP_PATH.read_bytes()
        form = asyncio.run(
            xlsx_extract(raw, "Cloud Security Principles Checklist V5_3.xlsx")
        )
        assert len(form.fields) == 45, (
            f"CSP field count drifted: {len(form.fields)} != 45 "
            f"(FORM-EXTRACTION-SPIKE.md §1.1 baseline)"
        )


class TestDR058GenericRuleIsolated:
    """Synthetic, isolated proofs that the shape-3 generic rule itself
    fires beyond the two archetypes — independent of the real corpus above,
    so a future change that narrows the rule fails here with a direct,
    minimal repro rather than only via an opaque real-fixture count drift."""

    def test_docx_generic_label_to_empty_cell_row_beyond_archetypes(self) -> None:
        """A 2-column table whose header does NOT classify as an authored
        Q/A table (``_classify_header`` returns None for both columns —
        neither archetype 1 nor archetype 2 can fire) still yields a field
        from a plain "label cell followed by an empty cell" row."""
        from docx import Document as _Document

        doc = _Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Company Registration Number"
        table.rows[0].cells[1].text = ""  # genuinely empty answer cell
        table.rows[1].cells[0].text = "VAT Number"
        table.rows[1].cells[1].text = ""
        import io

        buf = io.BytesIO()
        doc.save(buf)

        form = asyncio.run(docx_extract(buf.getvalue(), "synthetic-generic.docx"))
        labels = {f.question_text for f in form.fields if f.question_text}
        assert "Company Registration Number" in labels
        assert "VAT Number" in labels
        assert all(f.field_type == "empty_cell" for f in form.fields)

    def test_xlsx_generic_labelled_adjacent_empty_cell_beyond_archetypes(
        self,
    ) -> None:
        """A worksheet matching neither the CSP header signature
        (``Principle``/``Implementation``) nor the EFA signature
        (``Ref``/``Criteria``) still yields fields from repeated
        labelled-cell + empty-adjacent-cell rows (>= the 3-row repetition
        floor xlsx.py's generic fallback requires)."""
        import io

        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Supplier Details"
        ws["A1"] = "Field"
        ws["B1"] = "Response"
        ws["A2"] = "Company name"
        ws["A3"] = "Registered address"
        ws["A4"] = "Company registration number"
        # B2..B4 deliberately left blank — the unfilled answer cells.
        buf = io.BytesIO()
        wb.save(buf)

        form = asyncio.run(xlsx_extract(buf.getvalue(), "synthetic-generic.xlsx"))
        labels = {f.question_text for f in form.fields if f.question_text}
        assert "Company name" in labels
        assert "Registered address" in labels
        assert "Company registration number" in labels


class TestDR058GenericFallbackDoesNotFireOnMetadata:
    """Locks in the false-positive guard the header+repetition gating in
    ``xlsx.py`` exists for — measured against the EFA workbook's own
    non-scoring-matrix sheets during DR-058 development: a document-metadata
    block (scattered short label/value pairs with large gaps) and a
    sub-header row immediately following a real header must NOT produce
    fields on their own (below the repetition floor)."""

    def test_xlsx_sparse_metadata_block_yields_no_generic_fields(self) -> None:
        """A metadata block (title/owner/organisation rows with a gapped
        column layout) is NOT a repeating table — the generic fallback
        must require >= 3 repeating candidate rows before accepting a
        region, so an isolated 1-2-row coincidental match is discarded."""
        import io

        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Document Properties"
        ws["B24"] = "Document Properties"
        ws["B25"] = "Document Owner"
        ws["F25"] = "Deputy Director"
        ws["B26"] = "Organisation"
        ws["F26"] = "Acme Council"
        buf = io.BytesIO()
        wb.save(buf)

        form = asyncio.run(xlsx_extract(buf.getvalue(), "synthetic-metadata.xlsx"))
        assert form.fields == [], (
            "a sparse metadata block below the repetition floor must yield "
            "zero generic fields — a coincidental single-row match is not "
            "a genuine repeating fillable table"
        )


# ──────────────────────────────────────────────────────────────────────────
# DR-058 — w:sdt content controls + highlighted runs (TECH §3.2: "ALSO emit
# w:sdt content-controls + highlighted runs"). No real corpus fixture
# exercises either shape (Charnwood has 0 sdt elements and its 188
# highlighted runs are all non-empty buyer-customisation emphasis, not
# blank fill-in slots — see the module docstring on
# ``_highlighted_run_fields`` in docx.py for the measured false-positive
# evidence that shaped this design) — these are synthetic, built
# programmatically with raw OOXML elements (the same pattern
# ``test_docx_tracked_changes_regression.py`` uses for w:ins/w:del).
# ──────────────────────────────────────────────────────────────────────────


def _build_docx_with_sdt(
    *, alias: str | None, content_text: str | None
) -> bytes:
    """Build a minimal .docx whose body carries one ``w:sdt`` content
    control with an optional ``w:sdtPr/w:alias`` title and optional
    ``w:sdtContent`` text."""
    import io as _io

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    # Seed a real paragraph so the document is non-empty even if the sdt
    # sweep were somehow skipped.
    doc.add_paragraph("Cover sheet.")

    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    if alias is not None:
        alias_el = OxmlElement("w:alias")
        alias_el.set(qn("w:val"), alias)
        sdt_pr.append(alias_el)
    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")
    p = OxmlElement("w:p")
    if content_text is not None:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = content_text
        r.append(t)
        p.append(r)
    sdt_content.append(p)
    sdt.append(sdt_content)

    doc.element.body.insert(len(doc.element.body) - 1, sdt)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_with_highlighted_run(
    *, run_text: str, colour: str = "yellow"
) -> bytes:
    """Build a minimal .docx whose body carries one paragraph containing a
    single highlighted run (``w:rPr/w:highlight``)."""
    import io as _io

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(run_text)
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), colour)
    run._element.get_or_add_rPr().append(highlight)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDR058SdtContentControls:
    """``w:sdt`` content-control detection (synthetic — no real fixture)."""

    def test_sdt_with_alias_and_stock_placeholder_emits_field(self) -> None:
        raw = _build_docx_with_sdt(
            alias="Company Name", content_text="Click here to enter text."
        )
        form = asyncio.run(docx_extract(raw, "sdt-synthetic.docx"))
        matches = [f for f in form.fields if f.question_text == "Company Name"]
        assert matches, "expected a field carrying the sdt alias as question_text"
        field = matches[0]
        assert field.placeholder_text == "Click here to enter text."
        assert field.field_type == "placeholder"

    def test_sdt_with_alias_only_and_empty_content_emits_field(self) -> None:
        raw = _build_docx_with_sdt(alias="VAT Number", content_text=None)
        form = asyncio.run(docx_extract(raw, "sdt-synthetic.docx"))
        matches = [f for f in form.fields if f.question_text == "VAT Number"]
        assert matches, "an aliased, empty content control must still emit"
        assert matches[0].field_type == "empty_cell"
        assert matches[0].placeholder_text is None

    def test_sdt_already_filled_with_authored_content_emits_nothing(self) -> None:
        """A content control carrying real (non-stock) authored text is
        FILLED, not an unfilled slot — must not surface as a field."""
        raw = _build_docx_with_sdt(
            alias="Company Name", content_text="Acme Consulting Ltd"
        )
        form = asyncio.run(docx_extract(raw, "sdt-synthetic.docx"))
        assert not any(f.question_text == "Company Name" for f in form.fields)

    def test_sdt_with_no_alias_and_no_content_emits_nothing(self) -> None:
        """No label and no visible placeholder prompt — nothing to anchor
        a field on."""
        raw = _build_docx_with_sdt(alias=None, content_text=None)
        form = asyncio.run(docx_extract(raw, "sdt-synthetic.docx"))
        # The cover-sheet paragraph is the only expected field source (a
        # plain prose paragraph carries no placeholder span, so this is 0).
        assert form.fields == []


class TestDR058HighlightedRuns:
    """Highlighted-run detection (synthetic — no real fixture). Tightened
    to genuinely BLANK highlighted spans only after real-corpus measurement
    (Charnwood: 188 highlighted runs, all non-empty buyer-customisation
    emphasis — see docx.py's ``_highlighted_run_fields`` docstring)."""

    def test_empty_highlighted_run_emits_a_highlighted_field(self) -> None:
        raw = _build_docx_with_highlighted_run(run_text="")
        form = asyncio.run(docx_extract(raw, "highlight-synthetic.docx"))
        highlighted = [f for f in form.fields if f.field_type == "highlighted"]
        assert highlighted, "an empty highlighted run must emit a highlighted field"

    def test_nonempty_highlighted_prose_emits_nothing(self) -> None:
        """Highlighted PROSE is ambiguous emphasis, not reliably a fillable
        field (measured false-positive risk) — must not surface."""
        raw = _build_docx_with_highlighted_run(
            run_text="Charnwood Borough Council"
        )
        form = asyncio.run(docx_extract(raw, "highlight-synthetic.docx"))
        assert not any(f.field_type == "highlighted" for f in form.fields)

    def test_non_highlighted_run_emits_nothing(self) -> None:
        raw = _build_docx_with_highlighted_run(run_text="", colour="none")
        form = asyncio.run(docx_extract(raw, "highlight-synthetic.docx"))
        assert not any(f.field_type == "highlighted" for f in form.fields)
