"""Tracked-changes DOCX regression — ID-61.2.

Pre-reingest correctness gate for T7: the example-client corpus is .docx, and a
tracked-changes leak would corrupt extracted content on the headline
re-ingest run. ``scripts/docx_utils.py`` already resolves revisions
(``open_document_safe`` → pandoc ``--track-changes=accept``) and the
DOCX form extractor delegates to it — but until now only the Charnwood
acceptance fixture *incidentally* exercised that path. This module is
the dedicated regression: a fixture with genuine ``w:ins`` (insertion)
and ``w:del`` (deletion) revisions round-trips through
``form_extractors/docx.py`` and we assert, field by field, that

* deleted-run text is ABSENT from every extracted field;
* inserted-run text IS present as accepted content;
* no raw ``<w:ins>``/``<w:del>`` markup or revision-author metadata
  leaks into any field.

Fixture note — the .docx is built programmatically (python-docx +
OOXML revision elements) rather than committed as a binary blob: the
builder is reviewable in-diff, and ``TestFixtureCarriesTrackedChanges``
validates it against the PRODUCTION detector (``has_tracked_changes``)
so the fixture can never silently lose its revisions and let the
assertions pass vacuously.

Behavioural ground truth (verified against the real resolver):

* pandoc available — revisions are accepted before extraction:
  deletions discarded, insertions unwrapped into normal runs.
* pandoc absent — ``open_document_safe`` soft-warns and opens the raw
  file (the stance ``test_raises_when_no_pandoc`` pins for the helper).
  The document stays readable and no revision markup is injected, but
  insertions wrapped in ``w:ins`` may be missed by ``Paragraph.text``
  cell reads — exactly the degraded behaviour the runtime warning
  documents — so inserted-presence is asserted on the pandoc path only.
"""

from __future__ import annotations

import asyncio
import io
import logging
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.cocoindex_pipeline.form_extractors.docx import (
    extract as docx_extract,
)
from scripts.cocoindex_pipeline.form_extractors.shared import ExtractedForm
from scripts.docx_utils import (
    _check_pandoc_available,
    get_track_changes_stats,
    has_tracked_changes,
)

# ──────────────────────────────────────────────────────────────────────────
# Fixture content markers. Distinctive tokens so a leak anywhere in the
# extracted output is unambiguous in the assertion failure message.
# ──────────────────────────────────────────────────────────────────────────

REVISION_AUTHOR = "Redline Reviewer"

# Paragraph placeholder — the inserted span is a bracketed [Insert …]
# placeholder the extractor's inline matcher recognises; the deleted
# span is ALSO placeholder-shaped, so a deletion leak would surface as
# a spurious extra placeholder field (not just stray prose).
INSERTED_PARA_PLACEHOLDER = "[Insert organisational email address]"
DELETED_PARA_PLACEHOLDER = "[Insert DELETED-FAX-MARKER facsimile number]"

# Authored Q/A row — question cell carries a base run plus one inserted
# and one deleted revision run.
QUESTION_BASE_TEXT = "Describe your organisation's approach to "
INSERTED_TABLE_TEXT = "INSERTED-SAFEGUARDING obligations"
DELETED_TABLE_TEXT = "DELETED-OBSOLETE-CLAUSE"

PANDOC_AVAILABLE = _check_pandoc_available()
requires_pandoc = pytest.mark.skipif(
    not PANDOC_AVAILABLE,
    reason="pandoc not installed — accept-revisions path not exercisable",
)


def _revision_wrapper(
    wrapper_tag: str, text_tag: str, text: str, rev_id: int
) -> OxmlElement:
    """Build a ``w:ins``/``w:del`` revision element wrapping one run.

    Mirrors the OOXML shape Word emits for tracked changes: deletions
    carry their text in ``w:delText`` (not ``w:t``), insertions wrap a
    normal ``w:t`` run. Both carry ``w:author``/``w:date`` revision
    metadata — the leak assertions below check that neither the markup
    nor the author string reaches any extracted field.
    """
    wrapper = OxmlElement(wrapper_tag)
    wrapper.set(qn("w:id"), str(rev_id))
    wrapper.set(qn("w:author"), REVISION_AUTHOR)
    wrapper.set(qn("w:date"), "2026-01-15T09:00:00Z")
    run = OxmlElement("w:r")
    text_el = OxmlElement(text_tag)
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    wrapper.append(run)
    return wrapper


def _build_tracked_changes_docx() -> bytes:
    """A minimal blank form carrying unaccepted tracked changes.

    Two surfaces, each with BOTH an insertion and a deletion:

    * a prose paragraph whose placeholder span is inside ``w:ins`` and
      whose superseded placeholder is inside ``w:del``;
    * an authored Q/A table (header ``Question``/``Answer``) whose
      question cell carries inserted and deleted revision runs, with an
      empty answer cell so the row emits as ``field_type='empty_cell'``.
    """
    doc = Document()

    para = doc.add_paragraph("Contact e-mail: ")
    para._p.append(
        _revision_wrapper("w:ins", "w:t", INSERTED_PARA_PLACEHOLDER, 101)
    )
    para._p.append(
        _revision_wrapper("w:del", "w:delText", DELETED_PARA_PLACEHOLDER, 102)
    )

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Question"
    table.rows[0].cells[1].text = "Answer"
    question_cell = table.rows[1].cells[0]
    question_cell.text = QUESTION_BASE_TEXT
    question_para = question_cell.paragraphs[0]._p
    question_para.append(
        _revision_wrapper("w:ins", "w:t", INSERTED_TABLE_TEXT, 103)
    )
    question_para.append(
        _revision_wrapper("w:del", "w:delText", DELETED_TABLE_TEXT, 104)
    )
    # Answer cell intentionally left empty — authored question with an
    # empty response cell (Inv-9), the common blank-form shape.

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture(scope="module")
def tracked_docx_bytes() -> bytes:
    return _build_tracked_changes_docx()


def _all_field_texts(form: ExtractedForm) -> list[str]:
    """Every text surface an extracted field exposes downstream."""
    texts: list[str] = []
    for field in form.fields:
        for value in (
            field.question_text,
            field.placeholder_text,
            field.section_name,
        ):
            if value:
                texts.append(value)
    return texts


class TestFixtureCarriesTrackedChanges:
    """Guard the fixture itself — if the builder ever stops producing
    real revisions, the leak assertions below would pass vacuously."""

    def test_production_detector_sees_revisions(
        self, tracked_docx_bytes: bytes, tmp_path
    ) -> None:
        path = tmp_path / "tracked-changes-fixture.docx"
        path.write_bytes(tracked_docx_bytes)
        assert has_tracked_changes(str(path)), (
            "fixture builder produced a document the production "
            "has_tracked_changes detector does not flag — the regression "
            "below would be vacuous"
        )
        stats = get_track_changes_stats(str(path))
        assert stats["insertion_count"] > 0
        assert stats["deletion_count"] > 0


@requires_pandoc
class TestTrackedChangesResolvedOnExtract:
    """ID-61.2 acceptance — fixture with w:ins + w:del round-trips
    through form_extractors/docx.py with revisions accepted."""

    @pytest.fixture(scope="class")
    def tracked_form(self, tracked_docx_bytes: bytes) -> ExtractedForm:
        return asyncio.run(
            docx_extract(tracked_docx_bytes, "tracked-changes-fixture.docx")
        )

    def test_deleted_text_absent_from_all_fields(
        self, tracked_form: ExtractedForm
    ) -> None:
        """(a) Deleted-run text must not appear in any question_text /
        placeholder_text — a leak here corrupts re-ingested content."""
        for text in _all_field_texts(tracked_form):
            assert "DELETED-FAX-MARKER" not in text, (
                f"deleted paragraph placeholder leaked into a field: {text!r}"
            )
            assert DELETED_TABLE_TEXT not in text, (
                f"deleted table-cell text leaked into a field: {text!r}"
            )

    def test_deleted_placeholder_emits_no_spurious_field(
        self, tracked_form: ExtractedForm
    ) -> None:
        """The deleted span is itself placeholder-shaped — a deletion
        leak would emit a spurious extra placeholder field."""
        leaked = [
            f
            for f in tracked_form.fields
            if f.placeholder_text and "DELETED-FAX-MARKER" in f.placeholder_text
        ]
        assert leaked == [], (
            "deleted [Insert …] span surfaced as its own placeholder field"
        )

    def test_inserted_text_present_as_accepted_content(
        self, tracked_form: ExtractedForm
    ) -> None:
        """(b) Inserted-run text must survive as normal accepted content
        in BOTH surfaces — the prose placeholder and the Q/A cell."""
        placeholder_fields = [
            f for f in tracked_form.fields if f.placeholder_text
        ]
        assert any(
            f.placeholder_text == INSERTED_PARA_PLACEHOLDER
            for f in placeholder_fields
        ), (
            "inserted [Insert …] placeholder was not extracted from the "
            f"paragraph — placeholders seen: "
            f"{[f.placeholder_text for f in placeholder_fields]!r}"
        )

        question_fields = [f for f in tracked_form.fields if f.question_text]
        assert any(
            INSERTED_TABLE_TEXT in (f.question_text or "")
            and f.table_index is not None
            for f in question_fields
        ), (
            "inserted run text missing from the authored Q/A cell — "
            f"questions seen: "
            f"{[f.question_text for f in question_fields]!r}"
        )

    def test_no_revision_markup_or_author_metadata_leaks(
        self, tracked_form: ExtractedForm
    ) -> None:
        """(c) No raw OOXML revision markup or revision-author metadata
        may reach any extracted field."""
        assert tracked_form.fields, "extraction yielded zero fields"
        for text in _all_field_texts(tracked_form):
            for marker in ("w:ins", "w:del", "<w:", REVISION_AUTHOR):
                assert marker not in text, (
                    f"revision artefact {marker!r} leaked into field text: "
                    f"{text!r}"
                )


class TestTrackedChangesWithoutPandoc:
    """Warning path — pandoc unavailable. ``open_document_safe`` must
    soft-warn and leave the document readable without injecting any
    revision markup (matching ``test_raises_when_no_pandoc``'s stance
    for the underlying helper)."""

    def test_warning_path_readable_without_markup_injection(
        self, tracked_docx_bytes: bytes, caplog: pytest.LogCaptureFixture
    ) -> None:
        with patch(
            "scripts.docx_utils._check_pandoc_available", return_value=False
        ):
            with caplog.at_level(logging.WARNING, logger="scripts.docx_utils"):
                form = asyncio.run(
                    docx_extract(
                        tracked_docx_bytes, "tracked-changes-fixture.docx"
                    )
                )

        # The patch must actually have routed us down the warning path —
        # otherwise these assertions silently test the pandoc path.
        assert any(
            "pandoc is not installed" in record.getMessage()
            for record in caplog.records
        ), "no-pandoc warning did not fire — patch target drifted?"

        # Readable: the prose placeholder still extracts (its text nodes
        # are reachable), and no field carries revision markup, author
        # metadata, or deleted text (w:delText is never read as w:t).
        assert form.fields, "warning path produced zero fields"
        texts = _all_field_texts(form)
        assert any(INSERTED_PARA_PLACEHOLDER in t for t in texts)
        for text in texts:
            for marker in (
                "w:ins",
                "w:del",
                "<w:",
                REVISION_AUTHOR,
                "DELETED-FAX-MARKER",
                DELETED_TABLE_TEXT,
            ):
                assert marker not in text, (
                    f"revision artefact {marker!r} leaked on the no-pandoc "
                    f"warning path: {text!r}"
                )
