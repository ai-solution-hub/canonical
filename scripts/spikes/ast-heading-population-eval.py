#!/usr/bin/env python3
"""
ID-56 {56.18} research spike — AST heading-population go/no-go evidence harness.

NON-PRODUCTION. One-shot spike harness backing
docs/specs/id-56-content-model-invariants/56.18-ast-heading-spike.md.
Run end-to-end; the captured numbers live in that doc.

Three parts:
  Stage 1  — API introspection capture: what does cocoindex 1.0.3
             RecursiveSplitter(language='markdown').split actually surface?
  Stage 2  — prototype mapping: derive heading_text / heading_level /
             heading_path / parent_chunk_id (as parent position) for each
             splitter chunk from offsets + a source-side heading index.
             Mirrors legacy semantics (scripts/kb_pipeline/chunk.py:78-107).
  Stage 3  — structural proxy eval (substituted for recall@k — justification
             in the spike doc §Stage-3): heading-boundary alignment, section
             purity, and derived-heading truthfulness, budget-split
             (language=None, as built at scripts/cocoindex_pipeline/flow.py)
             vs heading-aware (language='markdown'), Variant-B byte budgets
             (2000/200/1000, Liam-ratified {56.5}), over two corpus shapes:
               (a) UK-procurement markdown (heading-rich), and
               (b) the {56.5} extracted binary tender corpus (heading-poor —
                   what the pipeline actually ingests from docx/pdf/xls).

Run:
  PYTHONUNBUFFERED=1 python3 scripts/spikes/ast-heading-population-eval.py
  (no network, no DB, no credentials — pure local computation.)
"""

from __future__ import annotations

import inspect
import os
import re
import sys
from dataclasses import dataclass

# --- repo-relative paths (script lives in scripts/spikes/) ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, "..", ".."))
TEMPLATES = os.path.join(REPO_ROOT, "docs", "testing", "test-data", "templates")

# Variant-B byte budgets — Liam-ratified {56.5}; mirrors flow.py constants.
CHUNK_SIZE_BYTES = 2000
CHUNK_OVERLAP_BYTES = 200
CHUNK_MIN_SIZE_BYTES = 1000


# ---------------------------------------------------------------------------
# Stage 1 — API introspection capture
# ---------------------------------------------------------------------------


def stage1_introspection() -> None:
    import cocoindex
    from cocoindex._internal import core as _core
    from cocoindex.ops.text import RecursiveSplitter, detect_code_language
    from cocoindex.resources.chunk import Chunk

    print("=" * 72)
    print("STAGE 1 — API introspection (installed cocoindex)")
    print("=" * 72)
    print(f"cocoindex version : {cocoindex.__version__}")
    print(f"split signature   : {inspect.signature(RecursiveSplitter.split)}")
    print(f"Chunk fields      : {[f for f in Chunk.__slots__]}")

    # Raw PyO3 chunk attribute surface (what the Rust engine exposes before
    # the Python _convert_chunk wrapper).
    raw = _core.RecursiveSplitter(custom_languages=[]).split(
        "# Title\n\nBody.\n\n## Section\n\nMore.", 20, None, None, "markdown"
    )
    attrs = [a for a in dir(raw[0]) if not a.startswith("_")]
    print(f"raw PyO3 chunk attrs ({len(attrs)}): {attrs}")
    print(f"detect_code_language('x.md') -> {detect_code_language(filename='x.md')}")

    # Silent fallback check: an unknown language does NOT raise.
    s = RecursiveSplitter()
    r = s.split("# A\n\nB.", 10, language="definitely_not_a_language_xyz")
    print(f"unknown language: no error, {len(r)} chunks (silent fallback)")
    print()


# ---------------------------------------------------------------------------
# Stage 2 — prototype mapping (offsets + source-side heading index)
# ---------------------------------------------------------------------------

ATX_HEADING_RE = re.compile(r"^(#{1,6})\s+(\S.*)$")


@dataclass(frozen=True)
class Heading:
    offset: int  # char offset of the heading line start
    level: int
    text: str


def build_heading_index(text: str) -> list[Heading]:
    """ATX headings with char offsets, fence-aware.

    Limitation (recorded in the spike doc): ATX only — setext headings
    (underlined with === / ---) are not indexed. tree-sitter-md recognises
    them for BOUNDARY placement, so a production mapping must add them.
    """
    headings: list[Heading] = []
    in_fence = False
    pos = 0
    for line in text.splitlines(keepends=True):
        stripped = line.strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            in_fence = not in_fence
        elif not in_fence:
            m = ATX_HEADING_RE.match(line)
            if m:
                headings.append(
                    Heading(offset=pos, level=len(m.group(1)), text=m.group(2).strip())
                )
        pos += len(line)
    return headings


@dataclass(frozen=True)
class DerivedColumns:
    """The four content_chunks heading columns, derived per chunk.

    parent_position is the POSITION of the parent chunk (the chunk that
    starts the nearest ancestor heading's section); in flow.py it becomes
    parent_chunk_id = uuid5(_KH_PIPELINE_DOC_NS, f"chunk:{rel_path}:{parent_position}")
    — fully deterministic, no second-pass UPDATE needed (unlike the legacy
    scripts/kb_pipeline/chunk.py:266-281 PATCH pass).
    """

    heading_text: str | None
    heading_level: int | None
    heading_path: list[str]
    parent_position: int | None


def derive_heading_columns(
    chunk_starts: list[int], headings: list[Heading]
) -> list[DerivedColumns]:
    """Map each chunk (by start char offset) to the four heading columns.

    Semantics mirror the legacy chunker (scripts/kb_pipeline/chunk.py:78-107):
      - governing heading = deepest heading at-or-before the chunk start
        (the heading "in effect" where the chunk begins);
      - heading_path = ancestor stack from root down to and including it;
      - parent = the chunk that begins the nearest ancestor (level <
        governing level) heading's section. Chunks before any heading are
        preamble: NULL / NULL / [] / NULL.
    """
    results: list[DerivedColumns] = []
    # First chunk position whose governing heading IS h (for parent lookup).
    first_chunk_of_heading: dict[int, int] = {}

    for position, start in enumerate(chunk_starts):
        # Heading stack in effect at `start`.
        stack: list[Heading] = []
        for h in headings:
            if h.offset > start:
                break
            while stack and stack[-1].level >= h.level:
                stack.pop()
            stack.append(h)
        if not stack:
            results.append(DerivedColumns(None, None, [], None))
            continue
        governing = stack[-1]
        first_chunk_of_heading.setdefault(governing.offset, position)
        parent_position: int | None = None
        for anc in reversed(stack[:-1]):
            if anc.level < governing.level:
                parent_position = first_chunk_of_heading.get(anc.offset)
                break
        results.append(
            DerivedColumns(
                heading_text=governing.text,
                heading_level=governing.level,
                heading_path=[h.text for h in stack],
                parent_position=parent_position,
            )
        )
    return results


# ---------------------------------------------------------------------------
# Stage 3 — structural proxy eval
# ---------------------------------------------------------------------------


def section_span(governing: Heading, headings: list[Heading], doc_len: int) -> tuple[int, int]:
    """Char span of the section the governing heading owns: from the heading
    to the next heading of the same-or-higher level (or EOF)."""
    end = doc_len
    for h in headings:
        if h.offset > governing.offset and h.level <= governing.level:
            end = h.offset
            break
    return governing.offset, end


def analyse_doc(label: str, text: str) -> list[dict]:
    from cocoindex.ops.text import RecursiveSplitter

    splitter = RecursiveSplitter()
    headings = build_heading_index(text)
    rows = []
    for lang in (None, "markdown"):
        chunks = splitter.split(
            text,
            CHUNK_SIZE_BYTES,
            chunk_overlap=CHUNK_OVERLAP_BYTES,
            min_chunk_size=CHUNK_MIN_SIZE_BYTES,
            language=lang,
        )
        bounds = [(c.start.char_offset, c.end.char_offset) for c in chunks]
        starts = [st for st, _ in bounds]
        derived = derive_heading_columns(starts, headings)

        start_set = set(starts)
        aligned = sum(1 for h in headings if h.offset in start_set)
        pure = sum(
            1
            for (st, en) in bounds
            if not any(st < h.offset < en for h in headings)
        )
        # Truthfulness: fraction of chunk chars inside the section the
        # derived heading_path claims the chunk belongs to.
        truthful_chars = 0
        total_chars = 0
        for (st, en), cols in zip(bounds, derived):
            total_chars += en - st
            if cols.heading_level is None:
                span = (0, headings[0].offset if headings else len(text))
            else:
                governing = next(
                    h
                    for h in headings
                    if h.text == cols.heading_text and h.offset <= st
                    and h.level == cols.heading_level
                )
                span = section_span(governing, headings, len(text))
            truthful_chars += max(0, min(en, span[1]) - max(st, span[0]))
        rows.append(
            {
                "doc": label,
                "lang": str(lang),
                "chars": len(text),
                "headings": len(headings),
                "chunks": len(chunks),
                "heading_aligned": f"{aligned}/{len(headings)}" if headings else "n/a",
                "pure_chunks": f"{pure}/{len(chunks)}",
                "truthfulness": (truthful_chars / total_chars) if total_chars else 1.0,
            }
        )
    return rows


# --- {56.5}-style extractors for the binary tender corpus (shape b) ---------


def extract_docx(path: str) -> str:
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(path: str) -> str:
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        return "\n".join((pg.extract_text() or "") for pg in pdf.pages)


def extract_xls(path: str) -> str:
    import xlrd

    wb = xlrd.open_workbook(path)
    parts = []
    for sh in wb.sheets():
        parts.append(f"## Sheet: {sh.name}")
        for r in range(sh.nrows):
            cells = [str(sh.cell_value(r, c)).strip() for c in range(sh.ncols)]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


MARKDOWN_CORPUS = [
    ("long-terms.md", "__tests__/fixtures/cocoindex-chunking/long-terms.md"),
    ("short-clause.md", "__tests__/fixtures/cocoindex-chunking/short-clause.md"),
    ("test-bid-resources.md", "docs/testing/uat/test-bid-resources.md"),
]

EXTRACTED_CORPUS = [
    ("ITT Services.docx", "itt-services-charnwood/ITT Services.docx", extract_docx),
    (
        "SQ PPN 03/24.pdf",
        "sq-standard-selection-questionnaire/standard-selection-questionnaire-ppn-03-24.pdf",
        extract_pdf,
    ),
    ("ITT Eval Matrix.xls", "itt-services-charnwood/ITT Evaluation Matrix.xls", extract_xls),
]


def main() -> None:
    stage1_introspection()

    print("=" * 72)
    print("STAGE 2 — prototype mapping demo (long-terms.md, markdown mode)")
    print("=" * 72)
    from cocoindex.ops.text import RecursiveSplitter

    with open(os.path.join(REPO_ROOT, MARKDOWN_CORPUS[0][1])) as f:
        demo_text = f.read()
    demo_headings = build_heading_index(demo_text)
    demo_chunks = RecursiveSplitter().split(
        demo_text,
        CHUNK_SIZE_BYTES,
        chunk_overlap=CHUNK_OVERLAP_BYTES,
        min_chunk_size=CHUNK_MIN_SIZE_BYTES,
        language="markdown",
    )
    demo_cols = derive_heading_columns(
        [c.start.char_offset for c in demo_chunks], demo_headings
    )
    for pos, (c, cols) in enumerate(zip(demo_chunks, demo_cols)):
        path = " > ".join(cols.heading_path) or "(preamble)"
        print(
            f"  chunk {pos}: [{c.start.char_offset:5d},{c.end.char_offset:5d}) "
            f"H{cols.heading_level} {cols.heading_text!r} | path: {path} | "
            f"parent_position: {cols.parent_position}"
        )
    print()

    print("=" * 72)
    print("STAGE 3 — structural proxy eval (Variant-B budgets 2000/200/1000)")
    print("=" * 72)
    all_rows: list[dict] = []
    for label, rel in MARKDOWN_CORPUS:
        with open(os.path.join(REPO_ROOT, rel)) as f:
            all_rows.extend(analyse_doc(f"[md] {label}", f.read()))
    for label, rel, extractor in EXTRACTED_CORPUS:
        path = os.path.join(TEMPLATES, rel)
        if not os.path.exists(path):
            print(f"  SKIP (missing fixture): {rel}")
            continue
        try:
            all_rows.extend(analyse_doc(f"[extracted] {label}", extractor(path)))
        except Exception as e:  # noqa: BLE001 — spike harness, report and move on
            print(f"  SKIP ({type(e).__name__}): {rel}: {e}")

    hdr = f"{'doc':32} {'lang':9} {'chars':>7} {'hdgs':>4} {'chunks':>6} {'hdg-aligned':>11} {'pure':>9} {'truthful':>8}"
    print(hdr)
    print("-" * len(hdr))
    for r in all_rows:
        print(
            f"{r['doc']:32} {r['lang']:9} {r['chars']:>7} {r['headings']:>4} "
            f"{r['chunks']:>6} {r['heading_aligned']:>11} {r['pure_chunks']:>9} "
            f"{r['truthfulness']:>8.1%}"
        )


if __name__ == "__main__":
    sys.exit(main())
