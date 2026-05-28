#!/usr/bin/env python3
"""
ID-56 {56.5} research spike — RecursiveSplitter chunk-config recall@k eval.

Compares three RecursiveSplitter chunk configurations (A/B/C) on a recall@k
retrieval eval over a genuine UK-procurement corpus, to produce a Liam-ratifiable
recommendation for the chunk_size / chunk_overlap / min_chunk_size values that the
{56.8} chunking stage will adopt.

NON-PRODUCTION. This is a one-shot research spike. Run end-to-end and capture the
real numbers into docs/research/id56-5-recursive-splitter-eval.md.

API (verified against installed cocoindex==1.0.3 at authoring time):
  - Splitter:  from cocoindex.ops.text import RecursiveSplitter
               RecursiveSplitter().split(text, chunk_size, *, chunk_overlap=..., min_chunk_size=...)
               -> list[Chunk]; Chunk has .text, .start, .end (TextPosition with char_offset/byte_offset).
               chunk_size is in BYTES (PRODUCT C-30 / RESEARCH V-11). min_chunk_size default = chunk_size/2.
  - Embedder:  from cocoindex.ops.litellm import LiteLLMEmbedder
               LiteLLMEmbedder('text-embedding-3-large', dimensions=1024).embed(text) -> np.ndarray (async).
               bare model id; litellm auto-routes text-embedding-* to OpenAI, NO 'openai/' prefix.

  NOTE: cocoindex.functions.SplitRecursively does NOT exist in 1.0.3 — do not cite it (V-1 trap).

Run:
  PYTHONUNBUFFERED=1 python3 scripts/spikes/recursive-splitter-eval.py
  (sandbox must be disabled — real OpenAI network calls + mmap.)
"""

from __future__ import annotations

import asyncio
import os
import re
import sys
import warnings
from dataclasses import dataclass

import numpy as np

warnings.filterwarnings("ignore")

# --- repo-relative paths (script lives in scripts/spikes/) ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, "..", ".."))
TEMPLATES = os.path.join(REPO_ROOT, "docs", "testing", "test-data", "templates")


def load_openai_key() -> None:
    """Load OPENAI_API_KEY from .env.local at repo root (per brief; key confirmed present)."""
    if os.environ.get("OPENAI_API_KEY"):
        return
    env_path = os.path.join(REPO_ROOT, ".env.local")
    if not os.path.exists(env_path):
        sys.exit("STOP: .env.local not found at repo root.")
    with open(env_path) as f:
        for line in f:
            line = line.strip()
            if line.startswith("OPENAI_API_KEY="):
                os.environ["OPENAI_API_KEY"] = line.split("=", 1)[1].strip().strip('"').strip("'")
                return
    sys.exit("STOP: OPENAI_API_KEY missing from .env.local.")


# ---------------------------------------------------------------------------
# Corpus extraction
# ---------------------------------------------------------------------------


def extract_docx(path: str) -> str:
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(path: str) -> str:
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        return "\n".join((pg.extract_text() or "") for pg in pdf.pages)


def extract_xlsx(path: str) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"## Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xls(path: str) -> str:
    import xlrd

    wb = xlrd.open_workbook(path)
    parts = []
    for sh in wb.sheets():
        parts.append(f"## Sheet: {sh.name}")
        for r in range(sh.nrows):
            cells = [str(sh.cell_value(r, c)).strip() for c in range(sh.ncols)]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# Each corpus doc: id, relative path, extractor.
CORPUS_SPEC = [
    ("itt-services", "itt-services-charnwood/ITT Services.docx", extract_docx, "docx prose"),
    ("rfp-annex2", "rfp-british-council/annex_2_supplier_response.docx", extract_docx, "docx prose"),
    ("sq-ppn0324", "sq-standard-selection-questionnaire/standard-selection-questionnaire-ppn-03-24.pdf", extract_pdf, "pdf prose"),
    ("rfp-pricing", "rfp-british-council/annex_3_pricing_approach.xlsx", extract_xlsx, "xlsx tabular"),
    ("efa-evalmatrix", "itt-services-efa/evaluation-matrix-itt-vol8.xlsx", extract_xlsx, "xlsx tabular"),
    ("csp-checklist", "csp-checklist/Cloud Security Principles Checklist V5_3.xlsx", extract_xlsx, "xlsx tabular"),
    ("itt-evalmatrix-xls", "itt-services-charnwood/ITT Evaluation Matrix.xls", extract_xls, "xls tabular"),
]


@dataclass
class Doc:
    id: str
    text: str
    fmt: str
    rel_path: str


def build_corpus() -> list[Doc]:
    docs: list[Doc] = []
    for doc_id, rel, fn, fmt in CORPUS_SPEC:
        path = os.path.join(TEMPLATES, rel)
        if not os.path.exists(path):
            print(f"  [skip] {doc_id}: file not found ({rel})")
            continue
        try:
            text = fn(path)
        except Exception as ex:  # noqa: BLE001
            print(f"  [skip] {doc_id}: extract failed {type(ex).__name__}: {ex}")
            continue
        if not text.strip():
            print(f"  [skip] {doc_id}: empty extract")
            continue
        docs.append(Doc(id=doc_id, text=text, fmt=fmt, rel_path=rel))
        print(f"  [ok]   {doc_id:22s} {fmt:14s} {len(text):>7d} chars")
    return docs


# ---------------------------------------------------------------------------
# Query set — authored by READING the docs (non-circular ground truth).
# Each query targets a doc and a distinctive answer SUBSTRING; the ground-truth
# span is located by .find() on the extracted text at runtime (per variant the
# chunk boundaries differ, so overlap is recomputed each time).
# ---------------------------------------------------------------------------

# (query_text, target_doc_id, answer_substring)
QUERIES = [
    ("What insurance cover does the contract require?", "itt-services", "Insurance"),
    ("What happens to a tender that is late or incomplete?", "itt-services", "No Tender will be considered which is late or incomplete"),
    ("How can a bidder seek clarification on the tender documents?", "itt-services", "You may seek clarification on any of the points contained in the Tender documents"),
    ("On what basis is each submission scored — price versus quality?", "itt-services", "scoring each submission on a"),
    ("Does the contract deal with TUPE and re-tendering?", "itt-services", "Tupe"),
    ("Is the Council liable for inaccuracy in the tender information?", "itt-services", "disclaim any liability for any inaccuracy or incompleteness"),
    ("Does the contract cover limitation of liability and indemnities?", "itt-services", "Limitation of liability"),
    ("What economic and financial standing checks apply to suppliers?", "sq-ppn0324", "economic and financial\nstanding"),
    ("What annual turnover threshold triggers a modern slavery statement?", "sq-ppn0324", "annual turnover of at least"),
    ("Can suppliers self-certify that they hold required insurance?", "sq-ppn0324", "self-certify that they have"),
    ("Which regulation provides statutory guidance on the selection stage?", "sq-ppn0324", "Regulation 107 of the PCR 2015"),
    ("What are the grounds for mandatory exclusion of a supplier?", "sq-ppn0324", "Grounds for mandatory exclusion"),
    ("What data protection legislation applies when processing personal data?", "sq-ppn0324", "Data Protection Legislation"),
    ("How should bidders complete the rate card and resources pricing tab?", "rfp-pricing", "Rate Card and Resources"),
    ("Where are British Council travel and subsistence rates found?", "rfp-pricing", "Travel & Subsistence"),
    ("Who sets the weightings for each area in the EFA evaluation matrix?", "efa-evalmatrix", "weightings for each area are provided"),
    ("What is the purpose of the EFA ITT scoring matrix?", "efa-evalmatrix", "ITT Scoring Matrix is provided to assist Framework Users"),
    ("What does the data-in-transit protection cloud security principle require?", "csp-checklist", "Data in transit protection"),
    ("Which cloud security principle covers asset protection and resilience?", "csp-checklist", "Asset protection and resilience"),
    ("Which cloud security principle covers identity and authentication?", "csp-checklist", "Identity and authentication"),
    ("How is the supplier response scored against evaluation criteria?", "rfp-annex2", "scored according to the methodology"),
    ("What experience must the supplier describe about the project team?", "rfp-annex2", "experience of researching international cultural and arts programmes"),
]


# ---------------------------------------------------------------------------
# Variants
# ---------------------------------------------------------------------------

# (label, chunk_size_bytes, chunk_overlap_bytes) ; min_chunk_size = default = chunk_size/2
VARIANTS = [
    ("A", 1000, 100),
    ("B", 2000, 200),
    ("C", 4000, 400),
]

K_VALUES = [1, 5, 10]


@dataclass
class ChunkRec:
    doc_id: str
    text: str
    char_start: int
    char_end: int


def split_corpus(docs: list[Doc], chunk_size: int, chunk_overlap: int) -> list[ChunkRec]:
    from cocoindex.ops.text import RecursiveSplitter

    splitter = RecursiveSplitter()
    out: list[ChunkRec] = []
    for d in docs:
        chunks = splitter.split(d.text, chunk_size, chunk_overlap=chunk_overlap)
        for c in chunks:
            out.append(
                ChunkRec(
                    doc_id=d.id,
                    text=c.text,
                    char_start=c.start.char_offset,
                    char_end=c.end.char_offset,
                )
            )
    return out


def ground_truth_span(docs_by_id: dict[str, Doc], target_doc: str, answer_sub: str) -> tuple[int, int]:
    """Locate the answer substring's char span in the target doc's extracted text."""
    text = docs_by_id[target_doc].text
    idx = text.find(answer_sub)
    if idx < 0:
        # tolerant fallback: collapse whitespace in both and match
        norm = re.sub(r"\s+", " ", text)
        norm_sub = re.sub(r"\s+", " ", answer_sub)
        nidx = norm.find(norm_sub)
        if nidx < 0:
            raise ValueError(f"answer substring not found in {target_doc}: {answer_sub!r}")
        # map normalised index back approximately by re-walking
        # (whitespace collapse only shrinks indices; use a coarse mapping)
        return _approx_span_from_norm(text, nidx, len(norm_sub))
    return idx, idx + len(answer_sub)


def _approx_span_from_norm(text: str, norm_start: int, norm_len: int) -> tuple[int, int]:
    """Map a normalised-text index back to a raw-text span (coarse but bounded)."""
    raw_i = 0
    norm_i = 0
    start_raw = None
    while raw_i < len(text) and norm_i <= norm_start + norm_len:
        ch = text[raw_i]
        is_ws = ch.isspace()
        # normalised collapses runs of whitespace to single space
        if is_ws:
            if norm_i > 0 and not text[raw_i - 1].isspace():
                norm_i += 1  # one space for the run
            if start_raw is None and norm_i >= norm_start:
                start_raw = raw_i
        else:
            if start_raw is None and norm_i >= norm_start:
                start_raw = raw_i
            norm_i += 1
        raw_i += 1
    if start_raw is None:
        start_raw = 0
    return start_raw, min(len(text), raw_i)


def overlaps(a_start: int, a_end: int, b_start: int, b_end: int) -> bool:
    return a_start < b_end and b_start < a_end


def cosine_rank(query_vec: np.ndarray, chunk_matrix: np.ndarray) -> np.ndarray:
    """Return chunk indices sorted by descending cosine similarity to query_vec.

    chunk_matrix rows are pre-L2-normalised; query_vec is normalised here.
    """
    q = query_vec / (np.linalg.norm(query_vec) + 1e-12)
    sims = chunk_matrix @ q
    return np.argsort(-sims)


async def embed_texts(embedder, texts: list[str]) -> np.ndarray:
    """Embed a list of strings via the async LiteLLMEmbedder.embed (one call per text).

    Asserts each vector length == 1024.
    """
    vecs = []
    # modest concurrency to keep it tidy and avoid rate spikes
    sem = asyncio.Semaphore(8)

    async def one(t: str) -> np.ndarray:
        async with sem:
            v = await embedder.embed(t if t.strip() else " ")
            arr = np.asarray(v, dtype=np.float32)
            assert arr.shape[0] == 1024, f"expected 1024-dim, got {arr.shape}"
            return arr

    results = await asyncio.gather(*(one(t) for t in texts))
    vecs = list(results)
    return np.vstack(vecs)


def l2_normalise(matrix: np.ndarray) -> np.ndarray:
    norms = np.linalg.norm(matrix, axis=1, keepdims=True)
    return matrix / (norms + 1e-12)


async def main() -> None:
    load_openai_key()
    from cocoindex.ops.litellm import LiteLLMEmbedder

    print("=== ID-56 {56.5} RecursiveSplitter recall@k spike ===")
    print("cocoindex==1.0.3 | embedding=text-embedding-3-large dim=1024 | chunk_size in BYTES")
    print("\n[1] Building corpus from docs/testing/test-data/templates/ ...")
    docs = build_corpus()
    if len(docs) < 3:
        sys.exit(f"STOP: only {len(docs)} usable docs extracted (<3). Corpus too thin to run.")
    docs_by_id = {d.id: d for d in docs}

    # Validate every query's ground-truth substring exists (fail fast on circular/typo).
    print("\n[2] Validating query ground-truth substrings ...")
    for i, (qtext, target, sub) in enumerate(QUERIES, 1):
        if target not in docs_by_id:
            sys.exit(f"STOP: query {i} targets missing doc {target}")
        try:
            gs, ge = ground_truth_span(docs_by_id, target, sub)
        except ValueError as ex:
            sys.exit(f"STOP: query {i} ground-truth not locatable: {ex}")
    print(f"  all {len(QUERIES)} queries grounded OK")

    embedder = LiteLLMEmbedder("text-embedding-3-large", dimensions=1024)

    # Embed queries once (variant-independent).
    print("\n[3] Embedding query set (one-off) ...")
    query_texts = [q[0] for q in QUERIES]
    query_matrix = await embed_texts(embedder, query_texts)
    print(f"  embedded {len(query_texts)} queries -> {query_matrix.shape}")

    results = {}  # label -> dict
    for label, csize, coverlap in VARIANTS:
        print(f"\n[4] Variant {label}: chunk_size={csize}B overlap={coverlap}B "
              f"min_chunk_size={csize // 2}B (default)")
        chunks = split_corpus(docs, csize, coverlap)
        chunk_texts = [c.text for c in chunks]
        sizes = [len(c.text.encode("utf-8")) for c in chunks]
        mean_bytes = sum(sizes) / len(sizes) if sizes else 0
        print(f"  {len(chunks)} chunks; mean {mean_bytes:.0f} bytes/chunk")

        print("  embedding chunks ...")
        chunk_matrix = l2_normalise(await embed_texts(embedder, chunk_texts))

        # Per-query recall.
        recall_hits = {k: 0 for k in K_VALUES}
        n_with_gt = 0
        for (qtext, target, sub), qvec in zip(QUERIES, query_matrix):
            gs, ge = ground_truth_span(docs_by_id, target, sub)
            # ground-truth chunk indices: chunks in target doc overlapping [gs,ge]
            gt_idx = {
                i for i, c in enumerate(chunks)
                if c.doc_id == target and overlaps(gs, ge, c.char_start, c.char_end)
            }
            if not gt_idx:
                # no chunk overlaps the answer span in this variant -> miss for all k
                continue
            n_with_gt += 1
            ranked = cosine_rank(qvec, chunk_matrix)
            for k in K_VALUES:
                topk = set(ranked[:k].tolist())
                if gt_idx & topk:
                    recall_hits[k] += 1

        total = len(QUERIES)
        recall = {k: recall_hits[k] / total for k in K_VALUES}
        results[label] = {
            "chunk_size": csize,
            "chunk_overlap": coverlap,
            "min_chunk_size": csize // 2,
            "n_chunks": len(chunks),
            "mean_bytes": mean_bytes,
            "n_with_gt": n_with_gt,
            "recall": recall,
        }
        print(f"  recall@1={recall[1]:.3f} recall@5={recall[5]:.3f} recall@10={recall[10]:.3f} "
              f"(queries with a GT chunk this variant: {n_with_gt}/{total})")

    # Final table.
    print("\n" + "=" * 70)
    print("RECALL@K SUMMARY (total queries = %d)" % len(QUERIES))
    print("=" * 70)
    hdr = f"{'Variant':8s} {'size/ovl/min(B)':16s} {'#chunks':8s} {'meanB':7s} " \
          f"{'r@1':6s} {'r@5':6s} {'r@10':6s}"
    print(hdr)
    print("-" * len(hdr))
    for label, _, _ in VARIANTS:
        r = results[label]
        cfg = f"{r['chunk_size']}/{r['chunk_overlap']}/{r['min_chunk_size']}"
        print(f"{label:8s} {cfg:16s} {r['n_chunks']:<8d} {r['mean_bytes']:<7.0f} "
              f"{r['recall'][1]:<6.3f} {r['recall'][5]:<6.3f} {r['recall'][10]:<6.3f}")
    print("=" * 70)
    print("\nDONE. Transcribe these REAL numbers into "
          "docs/research/id56-5-recursive-splitter-eval.md")


if __name__ == "__main__":
    asyncio.run(main())
