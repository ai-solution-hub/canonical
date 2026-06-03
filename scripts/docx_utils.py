"""Shared utilities for .docx file handling, including Track Changes resolution.

python-docx does NOT resolve tracked changes (revisions) in .docx files.
Documents with unaccepted Track Changes will have incorrect text extracted --
deleted text may be included and inserted text may be missed. This module
provides pandoc-based pre-processing to resolve Track Changes before extraction.

Usage:
    from docx_utils import open_document_safe, has_tracked_changes

    # Safe open -- resolves Track Changes automatically
    doc, temp_path = open_document_safe("input.docx")
    try:
        # ... use doc ...
    finally:
        if temp_path:
            os.unlink(temp_path)

    # Detection only
    if has_tracked_changes("input.docx"):
        print("Warning: file contains unresolved Track Changes")
"""

import logging
import os
import shutil
import subprocess
import tempfile

from docx import Document

logger = logging.getLogger(__name__)


def _resolve_pandoc_path() -> str | None:
    """Resolve a usable pandoc binary path.

    Prefers a system ``pandoc`` on PATH; falls back to the static binary
    bundled by the ``pypandoc_binary`` wheel. The cocoindex pipeline image is
    buildpack-built ({66.7}) with no apt layer, so pandoc ships as a
    pip-installed binary rather than a system package ({80.3}). Returns None
    when neither is available (the caller then soft-warns and skips resolution).
    """
    path = shutil.which("pandoc")
    if path is not None:
        return path
    try:
        import pypandoc  # noqa: PLC0415

        return pypandoc.get_pandoc_path()
    except (ImportError, OSError):
        return None


def _check_pandoc_available() -> bool:
    """Check if pandoc is available (system PATH or pypandoc-bundled binary)."""
    return _resolve_pandoc_path() is not None


def has_tracked_changes(file_path: str) -> bool:
    """Check if a .docx file contains unaccepted Track Changes.

    Inspects the OOXML body for <w:ins> (insertions) and <w:del> (deletions).
    This is a reliable detection method but does not reveal which specific
    cells or text spans are affected.

    Args:
        file_path: Path to the .docx file

    Returns:
        True if the document contains unresolved Track Changes
    """
    doc = Document(file_path)
    body_xml = doc.element.body.xml
    return "w:ins" in body_xml or "w:del" in body_xml


def get_track_changes_stats(file_path: str) -> dict:
    """Get detailed Track Changes statistics for a .docx file.

    Args:
        file_path: Path to the .docx file

    Returns:
        Dict with keys: has_changes, insertion_count, deletion_count
    """
    doc = Document(file_path)
    body_xml = doc.element.body.xml
    ins_count = body_xml.count("w:ins")
    del_count = body_xml.count("w:del")
    return {
        "has_changes": ins_count > 0 or del_count > 0,
        "insertion_count": ins_count,
        "deletion_count": del_count,
    }


def resolve_track_changes(input_path: str) -> str:
    """Accept all Track Changes in a .docx file using pandoc.

    Pandoc parses the OOXML at the revision level. With --track-changes=accept:
    - Text inside <w:del> elements is removed (deleted content discarded)
    - Text inside <w:ins> elements is unwrapped (inserted content becomes normal)
    - The result is a clean document with only the final/accepted text

    Args:
        input_path: Path to the .docx file with Track Changes

    Returns:
        Path to a cleaned temporary .docx file. The caller is responsible
        for cleaning up the temp file (e.g. os.unlink(path)).

    Raises:
        FileNotFoundError: If pandoc is not installed
        subprocess.CalledProcessError: If pandoc fails to process the file
    """
    pandoc_path = _resolve_pandoc_path()
    if pandoc_path is None:
        raise FileNotFoundError(
            "pandoc is not installed. Install with: brew install pandoc "
            "(or `pip install pypandoc_binary` for a bundled static binary)."
        )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        subprocess.run(
            [pandoc_path, "--track-changes=accept", "-o", tmp_path, input_path],
            check=True,
            capture_output=True,
            timeout=60,
        )
    except subprocess.CalledProcessError as e:
        # Clean up temp file on failure
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        logger.error(
            "pandoc failed to resolve Track Changes in %s: %s",
            os.path.basename(input_path),
            e.stderr.decode() if e.stderr else str(e),
        )
        raise

    return tmp_path


def open_document_safe(file_path: str) -> tuple[Document, str | None]:
    """Open a .docx file, resolving Track Changes if present.

    If the document contains unresolved Track Changes and pandoc is
    available, resolves them automatically. Otherwise opens the file
    directly (with a warning if Track Changes are present but pandoc
    is unavailable).

    Args:
        file_path: Path to the .docx file

    Returns:
        Tuple of (Document, temp_path_or_None). If temp_path is not None,
        the caller must delete it when done (os.unlink(temp_path)).
    """
    if has_tracked_changes(file_path):
        basename = os.path.basename(file_path)

        if _check_pandoc_available():
            logger.warning(
                "Track Changes detected in %s -- resolving via pandoc",
                basename,
            )
            clean_path = resolve_track_changes(file_path)
            return Document(clean_path), clean_path
        else:
            logger.warning(
                "Track Changes detected in %s but pandoc is not installed. "
                "Extracted text may include deleted content or miss inserted "
                "content. Install pandoc: brew install pandoc",
                basename,
            )

    return Document(file_path), None
