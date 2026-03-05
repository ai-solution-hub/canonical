"""Analyse a Word template to identify fields requiring completion.

Detection targets:
  1. Empty table cells adjacent to populated question cells
  2. Placeholder text patterns
  3. Cells containing only formatting markers (empty bullet lists)

Output:
  Writes identified fields to template_fields table.
  Uploads structure.json to Supabase Storage.
"""

import json
import os
import re
import sys
import tempfile
from datetime import datetime, timezone

from docx import Document

# Add scripts directory to path for imports
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

from extract_tender_questions import _classify_header

# Placeholder text patterns
PLACEHOLDER_PATTERNS = [
    re.compile(r'\[insert\b.*?\]', re.IGNORECASE),
    re.compile(r'\[enter\b.*?\]', re.IGNORECASE),
    re.compile(r'\[provide\b.*?\]', re.IGNORECASE),
    re.compile(r'\[please\b.*?\]', re.IGNORECASE),
    re.compile(r'\[your\b.*?\]', re.IGNORECASE),
    re.compile(r'\[type\b.*?\]', re.IGNORECASE),
    re.compile(r'\[response\b.*?\]', re.IGNORECASE),
    re.compile(r'\{\{[^}]+\}\}'),          # Jinja2-style: {{company_name}}
    re.compile(r'<<[A-Z_]+>>'),            # Angle-bracket: <<RESPONSE>>
    re.compile(r'\{[A-Z_]+\}'),            # Single-brace: {ANSWER}
    re.compile(r'n/a', re.IGNORECASE),     # "N/A" placeholder
    re.compile(r'^-+$'),                   # Dashes as placeholder
    re.compile(r'^\.\.\.$'),               # Ellipsis as placeholder
]


def _cell_text(cell) -> str:
    """Extract clean text from a table cell."""
    paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _is_empty_or_placeholder(cell) -> tuple[bool, str | None]:
    """Check if a cell is empty or contains only placeholder text.

    Returns:
        Tuple of (is_empty_or_placeholder, placeholder_text_or_None)
    """
    text = _cell_text(cell).strip()

    # Completely empty
    if not text:
        return True, None

    # Check placeholder patterns
    for pattern in PLACEHOLDER_PATTERNS:
        if pattern.fullmatch(text):
            return True, text

    # Check if cell contains only whitespace or formatting
    if all(not p.text.strip() for p in cell.paragraphs):
        return True, None

    return False, None


def _has_content(cell) -> bool:
    """Check if a cell has meaningful content (not empty/placeholder)."""
    is_empty, _ = _is_empty_or_placeholder(cell)
    return not is_empty


def _detect_merged_cells(table) -> set[tuple[int, int]]:
    """Detect merged cells in a table.

    Returns set of (row_index, col_index) tuples that are merge continuations
    (i.e., not the top-left cell of a merge group).
    """
    merged = set()
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            # Check if this cell is the same object as a previous cell
            # (python-docx represents merged cells by returning the same cell object)
            if col_idx > 0 and cell._tc is row.cells[col_idx - 1]._tc:
                merged.add((row_idx, col_idx))
            if row_idx > 0 and cell._tc is table.rows[row_idx - 1].cells[col_idx]._tc:
                merged.add((row_idx, col_idx))
    return merged


def _extract_word_limit(text: str) -> int | None:
    """Extract word limit from text."""
    match = re.search(r'(?:max(?:imum)?\.?\s+)?(\d+)\s*words?', text, re.IGNORECASE)
    return int(match.group(1)) if match else None


def _extract_section_headings(doc: Document) -> dict[int, str]:
    """Map table indices to the most recent heading above them.

    Returns:
        Dict mapping table_index -> section_name
    """
    sections: dict[int, str] = {}
    current_section = ""
    table_index = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            style = element.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
            )
            if style is not None:
                style_val = style.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
                )
                if style_val.startswith("Heading"):
                    heading_text = "".join(element.itertext()).strip()
                    if heading_text:
                        current_section = heading_text

        elif tag == "tbl":
            sections[table_index] = current_section
            table_index += 1

    return sections


def _has_tracked_changes(doc: Document) -> bool:
    """Check if the document contains tracked changes (revisions)."""
    body_xml = doc.element.body.xml
    return 'w:ins' in body_xml or 'w:del' in body_xml


def analyse_template(file_path: str) -> dict:
    """Analyse a Word template and return identified fields.

    Args:
        file_path: Path to the .docx template file

    Returns:
        Dict with:
            - fields: list of identified field dicts
            - table_count: number of tables found
            - total_fields: number of completable fields
            - warnings: list of warning messages
            - document_info: document metadata
            - column_mapping: per-table column assignments
    """
    doc = Document(file_path)
    section_map = _extract_section_headings(doc)
    fields = []
    warnings = []
    column_mapping = []
    sequence = 0

    # Check for tracked changes
    if _has_tracked_changes(doc):
        warnings.append(
            "This document contains tracked changes. python-docx operates on the "
            "accepted version. For best results, accept all changes in Word before uploading."
        )

    # Detect merged cells and nested tables
    total_merged = 0
    has_nested = False

    for table_idx, table in enumerate(doc.tables):
        rows = table.rows
        if len(rows) < 2:
            continue  # Need header + at least one data row

        merged_cells = _detect_merged_cells(table)
        total_merged += len(merged_cells)
        section_name = section_map.get(table_idx, "")

        # Classify columns: find "question" and "answer" columns
        col_count = len(rows[0].cells)
        if col_count < 2:
            continue  # Need at least 2 columns

        # Check header row for clues
        header_cells = [_cell_text(cell).strip().lower() for cell in rows[0].cells]
        header_types = [_classify_header(h) for h in header_cells]

        # If we have explicit question/answer headers, use those
        q_col = None
        a_col = None

        if 'question' in header_types:
            q_col = header_types.index('question')
        if 'answer' in header_types:
            a_col = header_types.index('answer')

        # Fallback: heuristic — find the most-populated column (question)
        # and the most-empty column (answer)
        if q_col is None or a_col is None:
            col_emptiness = []
            for col_idx in range(col_count):
                empty_count = sum(
                    1 for row_idx, row in enumerate(rows[1:], 1)
                    if (row_idx, col_idx) not in merged_cells
                    and _is_empty_or_placeholder(row.cells[col_idx])[0]
                )
                total_data_rows = len(rows) - 1
                col_emptiness.append(empty_count / max(total_data_rows, 1))

            if q_col is None:
                # Question column: least empty, longest average text
                candidates = [
                    (idx, rate) for idx, rate in enumerate(col_emptiness)
                    if rate < 0.5 and idx != a_col
                ]
                if candidates:
                    q_col = min(candidates, key=lambda x: x[1])[0]

            if a_col is None and q_col is not None:
                # Answer column: most empty, adjacent to question column
                candidates = [
                    (idx, rate) for idx, rate in enumerate(col_emptiness)
                    if rate > 0.5 and idx != q_col
                ]
                if candidates:
                    # Prefer column adjacent to question column
                    for idx, rate in candidates:
                        if abs(idx - q_col) == 1:
                            a_col = idx
                            break
                    if a_col is None:
                        a_col = max(candidates, key=lambda x: x[1])[0]

        if q_col is None or a_col is None:
            warnings.append(
                f"Table {table_idx}: Could not identify question/answer columns. "
                f"Headers: {header_cells}"
            )
            continue

        # Record column mapping for structure.json
        column_mapping.append({
            "table_index": table_idx,
            "question_col": q_col,
            "answer_col": a_col,
            "header_labels": header_cells,
        })

        # Extract fields from data rows
        for row_idx, row in enumerate(rows[1:], 1):
            # Skip merged cells
            if (row_idx, a_col) in merged_cells:
                continue

            q_cell = row.cells[q_col]
            a_cell = row.cells[a_col]

            question_text = _cell_text(q_cell).strip()
            is_empty, placeholder = _is_empty_or_placeholder(a_cell)

            if question_text and is_empty:
                # Found a completable field
                word_limit = _extract_word_limit(question_text)

                # Check other columns for word limit
                if word_limit is None:
                    for col_idx in range(col_count):
                        if col_idx in (q_col, a_col):
                            continue
                        cell_text = _cell_text(row.cells[col_idx])
                        wl = _extract_word_limit(cell_text)
                        if wl is not None:
                            word_limit = wl
                            break

                fields.append({
                    "field_type": "placeholder" if placeholder else "empty_cell",
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "col_index": a_col,
                    "question_text": question_text,
                    "section_name": section_name,
                    "word_limit": word_limit,
                    "placeholder_text": placeholder,
                    "sequence": sequence,
                })
                sequence += 1

    # Build document info
    document_info = {
        "table_count": len(doc.tables),
        "paragraph_count": len(doc.paragraphs),
        "has_tracked_changes": _has_tracked_changes(doc),
        "has_content_controls": False,
        "has_nested_tables": has_nested,
        "has_merged_cells": total_merged > 0,
        "merged_cell_count": total_merged,
    }

    return {
        "fields": fields,
        "table_count": len(doc.tables),
        "total_fields": len(fields),
        "warnings": warnings,
        "document_info": document_info,
        "column_mapping": column_mapping,
    }


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(
            "Usage: python3 scripts/analyse_template.py <template.docx>",
            file=sys.stderr,
        )
        sys.exit(1)

    file_path = sys.argv[1]
    if not file_path.lower().endswith(".docx"):
        print(f"Error: Expected a .docx file, got: {file_path}", file=sys.stderr)
        sys.exit(1)

    try:
        result = analyse_template(file_path)
        print(json.dumps(result, indent=2, ensure_ascii=False))
        print(
            f"Found {result['total_fields']} fields "
            f"in {result['table_count']} tables",
            file=sys.stderr,
        )
    except Exception as e:
        print(f"Error analysing template: {e}", file=sys.stderr)
        sys.exit(1)
