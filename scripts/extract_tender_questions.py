"""Extract tender questions from Word documents for bid response workflows.

Adapted from extract_docx_tables.py — reads tables from .docx files, identifies
question/requirement columns, captures word limits and section hierarchy, and
outputs structured JSON to stdout.

Handles multiple table formats:
  - Q&A pairs with Question/Response columns
  - Numbered question lists
  - Requirement matrices with evaluation criteria

Usage:
    python3 scripts/extract_tender_questions.py /path/to/tender.docx

Output:
    JSON to stdout with sections, questions, word limits, and sequence numbers.
    Errors and diagnostics go to stderr.
"""

import json
import os
import re
import sys
from typing import Optional

from docx import Document


# ── Header normalisation ────────────────────────────────────────────────

# Map common tender header text variants to canonical names
_QUESTION_HEADERS = {
    "question",
    "questions",
    "query",
    "requirement",
    "requirements",
    "response required",
    "response requirement",
    "description",
    "criteria",
    "criterion",
    "detail required",
    "details required",
    "information required",
    "submission requirement",
    "submission requirements",
    "tender question",
    "tender questions",
    "evaluation question",
    "evaluation questions",
    "question / requirement",
    "question/requirement",
    "scope",
    "specification",
}

_ANSWER_HEADERS = {
    "response",
    "answer",
    "reply",
    "supplier response",
    "tenderer response",
    "tenderer's response",
    "bidder response",
    "bidder's response",
    "contractor response",
    "contractor's response",
    "your response",
    "your answer",
    "proposed response",
    "standard response",
    "response text",
    "notes / response",
    "notes/response",
}

_SECTION_HEADERS = {
    "section",
    "category",
    "topic",
    "area",
    "lot",
    "part",
    "theme",
    "heading",
    "group",
}

_NUMBER_HEADERS = {
    "no",
    "no.",
    "#",
    "number",
    "ref",
    "ref.",
    "reference",
    "id",
    "q no",
    "q no.",
    "item",
    "item no",
    "item no.",
    "question no",
    "question no.",
    "question number",
}

_WEIGHT_HEADERS = {
    "weighting",
    "weight",
    "weight %",
    "weighting %",
    "evaluation weighting",
    "score weighting",
    "marks",
    "max marks",
    "max score",
    "maximum marks",
    "maximum score",
    "points",
    "max points",
    "percentage",
}

_WORD_LIMIT_HEADERS = {
    "word limit",
    "word count",
    "word count limit",
    "max words",
    "maximum words",
    "max word count",
    "character limit",
    "response length",
    "page limit",
}

# Administrative field patterns — questions matching these are "informational"
_ADMIN_PATTERNS = [
    re.compile(r"company\s+name", re.IGNORECASE),
    re.compile(r"organisation\s+name", re.IGNORECASE),
    re.compile(r"organization\s+name", re.IGNORECASE),
    re.compile(r"trading\s+name", re.IGNORECASE),
    re.compile(r"registered\s+address", re.IGNORECASE),
    re.compile(r"registration\s+number", re.IGNORECASE),
    re.compile(r"company\s+number", re.IGNORECASE),
    re.compile(r"vat\s+(registration\s+)?number", re.IGNORECASE),
    re.compile(r"contact\s+(name|email|telephone|phone|number)", re.IGNORECASE),
    re.compile(r"date\s+of\s+(incorporation|registration)", re.IGNORECASE),
    re.compile(r"duns\s+number", re.IGNORECASE),
    re.compile(r"(website|web\s+address|url)", re.IGNORECASE),
    re.compile(r"turnover|annual\s+revenue", re.IGNORECASE),
    re.compile(r"number\s+of\s+employees", re.IGNORECASE),
    re.compile(r"(parent|holding)\s+company", re.IGNORECASE),
    re.compile(r"sme\s+(status|classification|declaration)", re.IGNORECASE),
    re.compile(r"^name\s*$", re.IGNORECASE),
    re.compile(r"^address\s*$", re.IGNORECASE),
    re.compile(r"^postcode\s*$", re.IGNORECASE),
    re.compile(r"^email\s*(address)?\s*$", re.IGNORECASE),
    re.compile(r"^telephone\s*(number)?\s*$", re.IGNORECASE),
    re.compile(r"^signature\s*$", re.IGNORECASE),
]

# Word limit detection regex — matches "500 words", "Max 500 words", etc.
_WORD_LIMIT_RE = re.compile(
    r"(?:max(?:imum)?\.?\s+)?(\d+)\s*words?",
    re.IGNORECASE,
)

# Weight/percentage regex — matches "10%", "10 marks", "10 points"
_WEIGHT_RE = re.compile(
    r"(\d+(?:\.\d+)?)\s*(?:%|marks?|points?|score)",
    re.IGNORECASE,
)


def _normalise_header(text: str) -> str:
    """Clean and lowercase a header cell for matching."""
    cleaned = text.strip().lower()
    # Remove trailing punctuation
    cleaned = re.sub(r'[:\-_\.]+$', '', cleaned).strip()
    return cleaned


def _classify_header(text: str) -> Optional[str]:
    """Classify a header cell into a canonical type.

    Returns one of: "question", "answer", "section", "number", "weight",
    "word_limit", or None if unrecognised.
    """
    normalised = _normalise_header(text)
    if not normalised:
        return None

    if normalised in _QUESTION_HEADERS:
        return "question"
    if normalised in _ANSWER_HEADERS:
        return "answer"
    if normalised in _SECTION_HEADERS:
        return "section"
    if normalised in _NUMBER_HEADERS:
        return "number"
    if normalised in _WEIGHT_HEADERS:
        return "weight"
    if normalised in _WORD_LIMIT_HEADERS:
        return "word_limit"

    return None


def _cell_text(cell) -> str:
    """Extract clean text from a table cell, preserving paragraph breaks."""
    paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _is_admin_field(text: str) -> bool:
    """Check if question text matches an administrative field pattern."""
    for pattern in _ADMIN_PATTERNS:
        if pattern.search(text):
            return True
    return False


def _extract_word_limit(text: str) -> Optional[int]:
    """Extract word limit from text content.

    Looks for patterns like "Max 500 words", "500 words", "(max. 300 words)".
    Returns the integer word limit or None.
    """
    match = _WORD_LIMIT_RE.search(text)
    if match:
        return int(match.group(1))
    return None


def _extract_weight(text: str) -> Optional[float]:
    """Extract evaluation weight from text content.

    Looks for patterns like "10%", "10 marks", "10 points".
    Returns the numeric value or None.
    """
    match = _WEIGHT_RE.search(text)
    if match:
        return float(match.group(1))
    return None


# ── Table extraction ────────────────────────────────────────────────────

def _extract_questions_from_table(
    table,
    section_name: str,
    table_index: int,
) -> list[dict]:
    """Extract questions from a single Word table.

    Args:
        table: A python-docx Table object
        section_name: Current section name from document headings
        table_index: Index of this table in the document

    Returns:
        List of question dicts with keys: question_text, word_limit,
        evaluation_weight, category, section_name_from_col
    """
    rows = table.rows
    if len(rows) < 2:
        return []  # Need at least a header row and one data row

    # Extract and classify header row
    header_cells = [_cell_text(cell) for cell in rows[0].cells]
    header_types = [_classify_header(h) for h in header_cells]

    # Check we have at least a question column
    if "question" not in header_types:
        # Try positional fallback: if first column looks like questions
        # (long text), treat it as such
        if len(header_cells) >= 2:
            # Check if header row itself looks like a question (no real headers)
            all_unrecognised = all(ht is None for ht in header_types)
            if all_unrecognised:
                # No recognisable headers — skip this table
                return []
        return []

    q_idx = header_types.index("question")
    answer_idx = header_types.index("answer") if "answer" in header_types else None
    section_idx = header_types.index("section") if "section" in header_types else None
    weight_idx = header_types.index("weight") if "weight" in header_types else None
    word_limit_idx = (
        header_types.index("word_limit") if "word_limit" in header_types else None
    )

    questions = []
    for row in rows[1:]:
        cells = row.cells
        if len(cells) <= q_idx:
            continue

        question_text = _cell_text(cells[q_idx]).strip()
        if not question_text:
            continue

        # Determine category
        if _is_admin_field(question_text):
            category = "informational"
        else:
            category = "mandatory"

        # Extract word limit — check dedicated column first, then question text,
        # then neighbouring cells
        word_limit = None
        if word_limit_idx is not None and word_limit_idx < len(cells):
            wl_text = _cell_text(cells[word_limit_idx])
            word_limit = _extract_word_limit(wl_text)
            if word_limit is None and wl_text.strip().isdigit():
                word_limit = int(wl_text.strip())

        if word_limit is None:
            word_limit = _extract_word_limit(question_text)

        if word_limit is None:
            # Check answer column and adjacent cells for word limits
            for idx in range(len(cells)):
                if idx == q_idx:
                    continue
                cell_text = _cell_text(cells[idx])
                wl = _extract_word_limit(cell_text)
                if wl is not None:
                    word_limit = wl
                    break

        # Extract evaluation weight
        evaluation_weight = None
        if weight_idx is not None and weight_idx < len(cells):
            wt_text = _cell_text(cells[weight_idx])
            evaluation_weight = _extract_weight(wt_text)
            if evaluation_weight is None and wt_text.strip().replace(".", "").isdigit():
                try:
                    evaluation_weight = float(wt_text.strip())
                except ValueError:
                    pass

        # Extract section from column if present
        section_from_col = None
        if section_idx is not None and section_idx < len(cells):
            col_section = _cell_text(cells[section_idx]).strip()
            if col_section:
                section_from_col = col_section

        questions.append({
            "question_text": question_text,
            "word_limit": word_limit,
            "evaluation_weight": evaluation_weight,
            "category": category,
            "section_name_from_col": section_from_col,
        })

    return questions


# ── Main extraction ─────────────────────────────────────────────────────

def extract_tender_questions(file_path: str) -> dict:
    """Extract all tender questions from a Word document.

    Walks through the document, tracking section headings (Heading 1/2/3)
    to assign section names. Questions are grouped by section with
    monotonically increasing sequence numbers.

    Args:
        file_path: Path to the .docx file

    Returns:
        Dict with keys: sections, total_questions, total_sections
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)

    # Track sections and questions
    sections_map: dict[str, list[dict]] = {}  # section_name -> questions
    section_order: list[str] = []  # preserve insertion order
    current_section = ""
    table_index = 0

    # Iterate through document body elements in order to track headings
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Check if this paragraph is a heading
            style = element.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
            )
            if style is not None:
                style_val = style.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                    "",
                )
                if style_val.startswith("Heading"):
                    texts = element.itertext()
                    heading_text = "".join(texts).strip()
                    if heading_text:
                        current_section = heading_text

        elif tag == "tbl":
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                questions = _extract_questions_from_table(
                    table,
                    section_name=current_section,
                    table_index=table_index,
                )

                for q in questions:
                    # Determine section: prefer column value, then heading,
                    # then fallback to "General"
                    section_name = (
                        q.pop("section_name_from_col", None)
                        or current_section
                        or "General"
                    )

                    if section_name not in sections_map:
                        sections_map[section_name] = []
                        section_order.append(section_name)

                    sections_map[section_name].append(q)

            table_index += 1

    # Build output with monotonically increasing sequence numbers
    sections = []
    total_questions = 0

    for section_seq, section_name in enumerate(section_order):
        questions = sections_map[section_name]
        section_questions = []

        for q_seq, q in enumerate(questions):
            section_questions.append({
                "question_text": q["question_text"],
                "question_sequence": q_seq,
                "word_limit": q["word_limit"],
                "evaluation_weight": q["evaluation_weight"],
                "category": q["category"],
            })

        sections.append({
            "section_name": section_name,
            "section_sequence": section_seq,
            "questions": section_questions,
        })

        total_questions += len(section_questions)

    return {
        "sections": sections,
        "total_questions": total_questions,
        "total_sections": len(sections),
    }


# ── CLI entry point ──────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(
            "Usage: python3 scripts/extract_tender_questions.py <tender.docx>",
            file=sys.stderr,
        )
        sys.exit(1)

    file_path = sys.argv[1]

    if not file_path.lower().endswith(".docx"):
        print(f"Error: Expected a .docx file, got: {file_path}", file=sys.stderr)
        sys.exit(1)

    try:
        result = extract_tender_questions(file_path)
        print(json.dumps(result, indent=2, ensure_ascii=False))
        print(
            f"Extracted {result['total_questions']} questions "
            f"in {result['total_sections']} sections",
            file=sys.stderr,
        )
    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error extracting questions: {e}", file=sys.stderr)
        sys.exit(1)
