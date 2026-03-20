"""Extract Q&A pairs from Word documents containing bid library tables.

Supports three table formats detected by the client documentation:
  - Pattern A: 6-col audit format (No, Section, Question, Standard Response, Advanced Response, Notes)
  - Pattern B: 5-col DRAFT format (No, Section, Question, Standard Response, Notes)
  - Pattern C: 6-col numbered format (No, Question, Standard Answer, Advanced Answer, Section, Notes)

Usage:
    from scripts.extract_docx_tables import extract_qa_from_docx
    pairs = extract_qa_from_docx("/path/to/document.docx")

Each returned dict contains:
    - question_text: str
    - answer_standard: str
    - answer_advanced: str (empty if not present in format)
    - section_name: str
    - source_file: str
    - table_index: int
    - row_index: int
"""

import os
import re
from typing import Optional

from docx import Document
from docx_utils import open_document_safe


# ── Text deduplication ──────────────────────────────────────────────────

def deduplicate_repeated_text(text: str) -> str:
    """Remove repeated text runs from a string.

    When pandoc resolves Track Changes in .docx headings, the text content
    can be repeated 2-3x across XML text runs, e.g.
    "Product Support" -> "Product SupportProduct SupportProduct Support".

    This function detects when the entire string is a repeated pattern and
    returns just one copy. Only triggers when the string is an exact multiple
    of a substring (length > 2 chars).
    """
    if len(text) < 6:
        return text

    for length in range(len(text) // 2, 2, -1):
        prefix = text[:length]
        if len(text) % length == 0 and text == prefix * (len(text) // length):
            return prefix.strip()

    return text


# ── Header normalisation ────────────────────────────────────────────────

# Map common header text variants to canonical names
_HEADER_MAP = {
    # Question columns
    "question": "question",
    "questions": "question",
    "query": "question",
    "requirement": "question",
    "requirements": "question",
    "suggested questions": "question",
    # Standard response columns
    "standard response": "standard",
    "standard answer": "standard",
    "standard": "standard",
    "response": "standard",
    "answer": "standard",
    "answer for standard audit system": "standard",
    "answer for standard audits": "standard",
    "standard configuration answer": "standard",
    # Advanced response columns
    "advanced response": "advanced",
    "advanced answer": "advanced",
    "advanced": "advanced",
    "enhanced response": "advanced",
    "enhanced answer": "advanced",
    "answer for advanced audits": "advanced",
    "advanced audits answer": "advanced",
    # Section columns
    "section": "section",
    "category": "section",
    "topic": "section",
    "area": "section",
    # Number columns
    "no": "number",
    "no.": "number",
    "#": "number",
    "number": "number",
    "ref": "number",
    "id": "number",
    # Notes columns
    "notes": "notes",
    "comments": "notes",
    "note": "notes",
    "comment": "notes",
    # Standard Selection Questionnaire (PPN 03/24) columns
    # Part 1 — Supplier Information
    "supplier information": "section",
    "supplier name": "question",
    "contact name": "question",
    "contact details": "question",
    "registered address": "question",
    "company registration number": "question",
    "trading status": "question",
    "date of registration": "question",
    "sme": "question",
    "company size": "question",
    # Part 2 — Exclusion Grounds
    "exclusion grounds": "section",
    "grounds for mandatory exclusion": "section",
    "grounds for discretionary exclusion": "section",
    "mandatory exclusion": "section",
    "discretionary exclusion": "section",
    "self-cleaning": "question",
    # Part 3 — Selection Questions
    "selection questions": "section",
    "economic and financial standing": "section",
    "technical and professional ability": "section",
    "modern slavery": "section",
    "health and safety": "section",
    "environmental management": "section",
    "quality management": "section",
    "carbon reduction": "section",
    "steel": "section",
    "additional conditions of participation": "section",
    # Common SQ question-like headers
    "declaration": "question",
    "statement": "question",
    "evidence": "question",
    "details": "question",
    "description": "question",
    "please provide": "question",
    "please confirm": "question",
    "please describe": "question",
    "please state": "question",
    "please detail": "question",
    "please give details": "question",
    # Common SQ response-like headers
    "supplier response": "standard",
    "your response": "standard",
    "your answer": "standard",
    "tenderer response": "standard",
    "tenderer's response": "standard",
    "bidder response": "standard",
    "bidder's response": "standard",
    "contractor response": "standard",
    "applicant response": "standard",
    "applicant's response": "standard",
    "organisation response": "standard",
    # Guidance/instruction columns (treat as notes)
    "guidance": "notes",
    "guidance notes": "notes",
    "instructions": "notes",
    "max score": "notes",
    "weighting": "notes",
    "scoring": "notes",
    "max marks": "notes",
    "pass/fail": "notes",
}


def normalize_header(text: str) -> str:
    """Normalise a table header cell to a canonical name.

    Strips whitespace, lowercases, and maps known variants.
    Returns the canonical name or the cleaned text if no mapping found.
    """
    cleaned = text.strip().lower()
    # Remove trailing punctuation
    cleaned = re.sub(r'[:\-_]+$', '', cleaned).strip()
    return _HEADER_MAP.get(cleaned, cleaned)


# ── Table format detection ──────────────────────────────────────────────

def _infer_empty_headers(headers: list[str]) -> list[str]:
    """Infer canonical names for empty header columns.

    Some Audit template files have 'Question' in column 0 but empty strings
    in columns 1-2 that actually contain standard/advanced answers. This
    function fills in those blanks based on position.
    """
    normalised = [normalize_header(h) for h in headers]

    # Only apply if col 0 is "question" and we have empty columns after it
    if not normalised or normalised[0] != "question":
        return normalised

    # Find empty columns between question and known metadata columns
    empty_indices = [i for i in range(1, len(normalised)) if normalised[i] == ""]
    if not empty_indices:
        return normalised

    # Assign first empty as "standard", second as "advanced"
    result = list(normalised)
    if len(empty_indices) >= 1:
        result[empty_indices[0]] = "standard"
    if len(empty_indices) >= 2:
        result[empty_indices[1]] = "advanced"
    return result


def detect_table_format(headers: list[str]) -> Optional[str]:
    """Detect the table format from normalised header names.

    Returns:
        "audit_6col" — Pattern A: number, section, question, standard, advanced, notes
        "draft_5col" — Pattern B: number, section, question, standard, notes
        "numbered_6col" — Pattern C: number, question, standard, advanced, section, notes
        None — unrecognised format (not a Q&A table)
    """
    normalised = [normalize_header(h) for h in headers]

    has_question = "question" in normalised
    has_standard = "standard" in normalised
    has_advanced = "advanced" in normalised
    has_section = "section" in normalised
    has_number = "number" in normalised

    # Try inferring empty headers if we have question but no standard
    if has_question and not has_standard:
        normalised = _infer_empty_headers(headers)
        has_standard = "standard" in normalised
        has_advanced = "advanced" in normalised

    if not has_question or not has_standard:
        # Positional fallback: if all headers are empty or unrecognised,
        # guess layout from column count (matches original Track 4 script)
        all_empty = all(h.strip() == "" for h in headers)
        if all_empty or (not has_question and not has_standard):
            col_count = len(headers)
            if col_count == 5:
                return "positional_5col"
            elif col_count >= 6:
                return "positional_6col"
        return None

    col_count = len(normalised)

    if col_count >= 6 and has_advanced and has_section and has_number:
        # Distinguish Pattern A from Pattern C by column order
        q_idx = normalised.index("question")
        s_idx = normalised.index("section")
        if s_idx < q_idx:
            return "audit_6col"      # Pattern A: section before question
        else:
            return "numbered_6col"   # Pattern C: section after question

    if col_count >= 5 and has_section and has_number and not has_advanced:
        return "draft_5col"          # Pattern B: no advanced column

    # Fallback: if we have question + standard, treat as a generic table
    if has_question and has_standard:
        if has_advanced:
            return "audit_6col"
        return "draft_5col"

    return None


# ── Q&A extraction from a single table ──────────────────────────────────

def _cell_text(cell) -> str:
    """Extract clean text from a table cell, preserving paragraph breaks."""
    paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def extract_qa_from_table(table, section_name: str = "", table_index: int = 0,
                          source_file: str = "") -> list[dict]:
    """Extract Q&A pairs from a single Word table.

    Args:
        table: A python-docx Table object
        section_name: Current section name from document headings
        table_index: Index of this table in the document (for traceability)
        source_file: Source filename for provenance

    Returns:
        List of dicts with keys: question_text, answer_standard, answer_advanced,
        section_name, source_file, table_index, row_index
    """
    rows = table.rows
    if len(rows) < 2:
        return []  # Need at least a header row and one data row

    # Extract header row
    header_cells = [_cell_text(cell) for cell in rows[0].cells]

    fmt = detect_table_format(header_cells)
    if fmt is None:
        return []  # Not a Q&A table

    # Positional formats: row 0 is data, not a header
    data_start = 1
    if fmt == "positional_5col":
        q_idx = 0
        std_idx = 1
        adv_idx = None
        sec_idx = None
        data_start = 0  # First row is data
    elif fmt == "positional_6col":
        q_idx = 0
        std_idx = 1
        adv_idx = 2
        sec_idx = None
        data_start = 0  # First row is data
    else:
        # Use inferred headers (fills in empty columns) for column mapping
        normalised_headers = _infer_empty_headers(header_cells)

        # Build column index map
        col_map = {}
        for idx, name in enumerate(normalised_headers):
            if name not in col_map:
                col_map[name] = idx

        q_idx = col_map.get("question")
        std_idx = col_map.get("standard")
        adv_idx = col_map.get("advanced")
        sec_idx = col_map.get("section")

    if q_idx is None or std_idx is None:
        return []

    pairs = []
    for row_num, row in enumerate(rows[data_start:], start=data_start):
        cells = row.cells

        # Guard against rows shorter than expected
        if len(cells) <= max(q_idx, std_idx):
            continue

        question = _cell_text(cells[q_idx])
        standard = _cell_text(cells[std_idx])

        # Skip empty question rows (spacing/formatting rows)
        if not question.strip():
            continue

        advanced = ""
        if adv_idx is not None and adv_idx < len(cells):
            advanced = _cell_text(cells[adv_idx])

        # Use section from column if present, otherwise fall back to heading-based section
        row_section = section_name
        if sec_idx is not None and sec_idx < len(cells):
            cell_section = _cell_text(cells[sec_idx])
            if cell_section.strip():
                row_section = cell_section.strip()

        pairs.append({
            "question_text": question,
            "answer_standard": standard,
            "answer_advanced": advanced,
            "section_name": row_section,
            "source_file": source_file,
            "table_index": table_index,
            "row_index": row_num,
        })

    return pairs


# ── Main extraction from a DOCX file ────────────────────────────────────

def extract_qa_from_docx(file_path: str) -> list[dict]:
    """Extract all Q&A pairs from a Word document.

    Walks through the document, tracking section headings (Heading 1/2/3)
    to assign section names to Q&A pairs extracted from tables.

    Args:
        file_path: Path to the .docx file

    Returns:
        List of dicts, each containing a Q&A pair with metadata
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc, temp_path = open_document_safe(file_path)
    try:
        source_file = os.path.basename(file_path)

        all_pairs = []
        current_section = ""
        table_index = 0

        # Iterate through the document body elements in order to track
        # headings that appear before each table
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Check if this paragraph is a heading
                style = element.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
                )
                if style is not None:
                    style_val = style.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
                    )
                    if style_val.startswith("Heading"):
                        # Extract the text content of this heading
                        texts = element.itertext()
                        heading_text = "".join(texts).strip()
                        # Clean up duplicated text from merged/complex formatting runs.
                        # Take only the first line if the heading spans multiple lines
                        # (subsequent lines are often formatting artefacts).
                        if "\n" in heading_text:
                            heading_text = heading_text.split("\n")[0].strip()
                        # Deduplicate repeated text runs (pandoc Track Changes artefact).
                        heading_text = deduplicate_repeated_text(heading_text)
                        if heading_text:
                            current_section = heading_text

            elif tag == "tbl":
                # This is a table element — find the corresponding python-docx Table
                if table_index < len(doc.tables):
                    table = doc.tables[table_index]
                    pairs = extract_qa_from_table(
                        table,
                        section_name=current_section,
                        table_index=table_index,
                        source_file=source_file,
                    )
                    all_pairs.extend(pairs)
                table_index += 1

        return all_pairs
    finally:
        if temp_path:
            os.unlink(temp_path)


# ── CLI entry point ──────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python3 scripts/extract_docx_tables.py <file.docx> [file2.docx ...]")
        sys.exit(1)

    total = 0
    for path in sys.argv[1:]:
        pairs = extract_qa_from_docx(path)
        print(f"\n{os.path.basename(path)}: {len(pairs)} Q&A pairs")
        for p in pairs[:3]:
            print(f"  [{p['section_name']}] Q: {p['question_text'][:80]}")
            print(f"    A (std): {p['answer_standard'][:80]}")
            if p['answer_advanced']:
                print(f"    A (adv): {p['answer_advanced'][:80]}")
        if len(pairs) > 3:
            print(f"  ... and {len(pairs) - 3} more")
        total += len(pairs)

    print(f"\nTotal: {total} Q&A pairs from {len(sys.argv) - 1} files")
