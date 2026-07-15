"""Fill a Word template with approved bid responses.

Reads the original template, writes responses into identified cells,
preserves formatting, and saves as a new file.

CRITICAL: Never use cell.text = "value" -- this destroys all formatting.
Always use the run-level API (paragraph.clear() + paragraph.add_run()).
"""

import io
import json
import os
import sys

import openpyxl
from docx import Document
from pypdf import PdfReader, PdfWriter

from docx_utils import open_document_safe


def _cell_text(cell) -> str:
    """Extract clean text from a table cell."""
    paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _enforce_word_limit(text: str, word_limit: int | None) -> tuple[str, bool]:
    """Truncate text to word limit if necessary.

    Returns:
        Tuple of (text, was_truncated)
    """
    if word_limit is None:
        return text, False

    words = text.split()
    if len(words) <= word_limit:
        return text, False

    truncated = " ".join(words[:word_limit])
    return truncated, True


def _copy_cell_formatting(source_cell, target_cell, text: str):
    """Write text into target_cell, copying paragraph formatting from source_cell.

    Preserves:
      - Font name, size, bold, italic, colour
      - Paragraph alignment
      - Paragraph spacing (before/after)

    Does NOT handle:
      - Bullet lists within cells
      - Images within cells
    """
    # Get reference formatting from first populated paragraph in source
    ref_para = None
    ref_run = None
    for p in source_cell.paragraphs:
        if p.text.strip() and p.runs:
            ref_para = p
            ref_run = p.runs[0]
            break

    # Clear target cell
    target_para = target_cell.paragraphs[0]
    target_para.clear()

    # Remove any extra paragraphs (placeholder text may have added them)
    while len(target_cell.paragraphs) > 1:
        p = target_cell.paragraphs[-1]
        p._element.getparent().remove(p._element)

    # Handle multi-paragraph text (split on double newlines)
    paragraphs_text = text.split('\n\n')

    # Track the last inserted element so subsequent paragraphs are appended
    # in order (addnext inserts immediately after the reference element).
    last_inserted = target_para._element

    for i, para_text in enumerate(paragraphs_text):
        if i == 0:
            # Use existing first paragraph
            para = target_para
        else:
            # Add new paragraph after the last inserted one
            from docx.oxml.ns import qn
            new_p = last_inserted.makeelement(qn('w:p'), {})
            last_inserted.addnext(new_p)
            last_inserted = new_p
            from docx.text.paragraph import Paragraph
            para = Paragraph(new_p, target_cell)

        # Write text as a run
        run = para.add_run(para_text.strip())

        # Apply formatting from reference
        if ref_run:
            if ref_run.font.name:
                run.font.name = ref_run.font.name
            if ref_run.font.size:
                run.font.size = ref_run.font.size
            if ref_run.font.bold is not None:
                run.font.bold = ref_run.font.bold
            if ref_run.font.italic is not None:
                run.font.italic = ref_run.font.italic
            if ref_run.font.color and ref_run.font.color.rgb:
                run.font.color.rgb = ref_run.font.color.rgb

        if ref_para and i == 0:
            if ref_para.alignment is not None:
                para.alignment = ref_para.alignment
            if ref_para.paragraph_format.space_before:
                para.paragraph_format.space_before = ref_para.paragraph_format.space_before
            if ref_para.paragraph_format.space_after:
                para.paragraph_format.space_after = ref_para.paragraph_format.space_after


def fill_template(
    template_path: str,
    output_path: str,
    field_mappings: list[dict],
) -> dict:
    """Fill identified cells in a Word template with response text.

    Args:
        template_path: Path to the original .docx template
        output_path: Path to save the completed .docx
        field_mappings: List of dicts with:
            - table_index (int)
            - row_index (int)
            - col_index (int)
            - response_text (str)
            - word_limit (int | None)

    Returns:
        Dict with fields_filled, fields_skipped, fields_failed, truncated, errors
    """
    doc, temp_path = open_document_safe(template_path)

    try:
        return _fill_template_inner(doc, output_path, field_mappings)
    finally:
        if temp_path:
            os.unlink(temp_path)


def _fill_template_inner(
    doc,
    output_path: str,
    field_mappings: list[dict],
) -> dict:
    """Inner implementation of fill_template (extracted for try/finally cleanup)."""
    filled = 0
    skipped = 0
    failed = 0
    truncated_fields = []
    errors = []

    for mapping in field_mappings:
        table_idx = mapping["table_index"]
        row_idx = mapping["row_index"]
        col_idx = mapping["col_index"]
        response_text = mapping["response_text"]
        word_limit = mapping.get("word_limit")

        try:
            # Validate indices
            if table_idx >= len(doc.tables):
                errors.append({
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "error": f"Table index {table_idx} out of range (document has {len(doc.tables)} tables)",
                })
                failed += 1
                continue

            table = doc.tables[table_idx]

            if row_idx >= len(table.rows):
                errors.append({
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "error": f"Row index {row_idx} out of range (table has {len(table.rows)} rows)",
                })
                failed += 1
                continue

            row = table.rows[row_idx]

            if col_idx >= len(row.cells):
                errors.append({
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "col_index": col_idx,
                    "error": f"Column index {col_idx} out of range",
                })
                failed += 1
                continue

            target_cell = row.cells[col_idx]

            # Skip if response is empty
            if not response_text or not response_text.strip():
                skipped += 1
                continue

            # Enforce word limit
            text_to_write, was_truncated = _enforce_word_limit(
                response_text, word_limit
            )
            if was_truncated:
                truncated_fields.append({
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "original_words": len(response_text.split()),
                    "limit": word_limit,
                })

            # Find reference cell for formatting (question cell in same row)
            ref_cell = None
            for c_idx, cell in enumerate(row.cells):
                if c_idx != col_idx and _cell_text(cell).strip():
                    ref_cell = cell
                    break

            # Write response with formatting preservation
            _copy_cell_formatting(
                ref_cell or target_cell,
                target_cell,
                text_to_write,
            )

            filled += 1

        except Exception as e:
            errors.append({
                "table_index": table_idx,
                "row_index": row_idx,
                "error": str(e),
            })
            failed += 1

    # Save completed document
    doc.save(output_path)

    return {
        "fields_filled": filled,
        "fields_skipped": skipped,
        "fields_failed": failed,
        "truncated": truncated_fields,
        "errors": errors,
    }


def _validate_completed_document(original_path: str, completed_path: str) -> list[str]:
    """Compare original and completed documents for formatting issues."""
    original, orig_temp = open_document_safe(original_path)
    completed, comp_temp = open_document_safe(completed_path)
    try:
        warnings = []

        if len(original.tables) != len(completed.tables):
            warnings.append(
                f"Table count mismatch: original has {len(original.tables)}, "
                f"completed has {len(completed.tables)}"
            )

        for idx, (orig_table, comp_table) in enumerate(
            zip(original.tables, completed.tables)
        ):
            if len(orig_table.rows) != len(comp_table.rows):
                warnings.append(
                    f"Table {idx}: row count mismatch "
                    f"({len(orig_table.rows)} vs {len(comp_table.rows)})"
                )

        for idx, (orig_section, comp_section) in enumerate(
            zip(original.sections, completed.sections)
        ):
            if orig_section.page_width != comp_section.page_width:
                warnings.append(f"Section {idx}: page width changed")
            if orig_section.page_height != comp_section.page_height:
                warnings.append(f"Section {idx}: page height changed")

        return warnings
    finally:
        if orig_temp:
            os.unlink(orig_temp)
        if comp_temp:
            os.unlink(comp_temp)


# ══════════════════════════════════════════════════════════════════════════
# PDF WRITER — pypdf AcroForm value writes into the plane-2 fillable
# artefact (ID-145 {145.15}, TECH.md §3.3/§3.4).
# ══════════════════════════════════════════════════════════════════════════


def _pdf_field_names_by_sequence(pdf_bytes: bytes) -> dict[int, str]:
    """Rebuild the ``sequence -> AcroForm field-name`` map for a fillable
    PDF's Widget annotations.

    ADDRESSING GAP: ``form_instance_fields`` carries no ``field_name``
    column for PDF rows — {145.11}'s ``ExtractedField`` shape-adapter
    (``orchestrator._pdf_result_to_extracted_form``) deliberately drops the
    commonforms-assigned field name at the DB-write boundary (the
    GEOMETRY-PERSISTENCE decision documented there), leaving only
    ``table_index`` (repurposed as the 0-indexed page) and ``row_index``
    (repurposed as a document-wide reading-order ``sequence``). This
    function reconstructs the mapping by walking the SAME bytes in the
    IDENTICAL page-major + Widget-annotation order
    ``form_extractors/pdf.py``'s ``detect_pdf_fields`` used to assign that
    ``sequence`` in the first place — deterministic (no re-detection, just
    re-reading already-written AcroForm structure) and verified stable
    across a pypdf write/read round-trip (a re-fill pass's base artefact is
    a PRIOR pypdf-written completion, not always the original commonforms
    output).
    """
    reader = PdfReader(io.BytesIO(pdf_bytes))
    names: dict[int, str] = {}
    sequence = 0
    for page in reader.pages:
        for annotation in page.get("/Annots") or []:
            obj = annotation.get_object()
            if obj.get("/Subtype") != "/Widget":
                continue
            if obj.get("/FT") not in ("/Tx", "/Btn"):
                continue
            names[sequence] = str(obj.get("/T"))
            sequence += 1
    return names


def fill_pdf_template(
    input_path: str,
    output_path: str,
    field_mappings: list[dict],
) -> dict:
    """Write approved bid responses into a PDF's AcroForm fields.

    Args:
        input_path: Path to the base fillable PDF — the plane-2 fillable
            artefact ({145.11}'s ``{form_id}/fillable.pdf``) on a first
            pass, or a PRIOR ``template_completions`` output on a re-fill
            pass (BI-22 re-entrancy — see ``bid_worker.fill_template_job``).
        output_path: Path to save the filled PDF.
        field_mappings: List of dicts with:
            - table_index (int | None) — PDF's 0-indexed page (unused for
              addressing; carried through into error/truncation records
              only).
            - row_index (int) — the PDF's ``sequence`` (see
              ``_pdf_field_names_by_sequence``).
            - response_text (str)
            - word_limit (int | None)

    Returns:
        Dict with fields_filled, fields_skipped, fields_failed, truncated, errors.
    """
    with open(input_path, "rb") as f:
        pdf_bytes = f.read()

    field_names_by_sequence = _pdf_field_names_by_sequence(pdf_bytes)

    writer = PdfWriter(clone_from=io.BytesIO(pdf_bytes))

    filled = 0
    skipped = 0
    failed = 0
    truncated_fields: list[dict] = []
    errors: list[dict] = []
    values: dict[str, str] = {}

    for mapping in field_mappings:
        table_idx = mapping.get("table_index")
        sequence = mapping.get("row_index")
        response_text = mapping.get("response_text")
        word_limit = mapping.get("word_limit")

        field_name = field_names_by_sequence.get(sequence)
        if field_name is None:
            errors.append({
                "table_index": table_idx,
                "row_index": sequence,
                "error": f"No AcroForm widget at sequence {sequence} in the fillable artefact",
            })
            failed += 1
            continue

        if not response_text or not response_text.strip():
            skipped += 1
            continue

        text_to_write, was_truncated = _enforce_word_limit(
            response_text, word_limit
        )
        if was_truncated:
            truncated_fields.append({
                "table_index": table_idx,
                "row_index": sequence,
                "original_words": len(response_text.split()),
                "limit": word_limit,
            })

        values[field_name] = text_to_write
        filled += 1

    if values:
        writer.update_page_form_field_values(None, values)

    with open(output_path, "wb") as f:
        writer.write(f)

    return {
        "fields_filled": filled,
        "fields_skipped": skipped,
        "fields_failed": failed,
        "truncated": truncated_fields,
        "errors": errors,
    }


# ══════════════════════════════════════════════════════════════════════════
# XLSX WRITER — openpyxl cell writes at the recorded coords (ID-145
# {145.15}, TECH.md §3.3/§3.4).
# ══════════════════════════════════════════════════════════════════════════


def _xlsx_sheet_by_table_index(raw_bytes: bytes) -> dict[int, str]:
    """Rebuild the ``table_index -> sheet_name`` boundary map for an XLSX
    workbook.

    ADDRESSING GAP: ``form_extractors/xlsx.py``'s ``extract()`` increments
    ``table_index`` as a WORKBOOK-WIDE running counter across every sheet
    (not a sheet index), and neither ``ExtractedField`` nor
    ``form_instance_fields`` carries a ``sheet_name``/``sheet_index``
    column — a real gap for multi-sheet workbooks (e.g. the EFA archetype's
    multiple "Bidder N" sheets), where ``(row_index, col_index)`` alone is
    ambiguous. This function reconstructs the boundary by re-running the
    SAME per-sheet walk (``_walk_sheet``, imported — a deterministic,
    rule-based walk, not an ML detector) over the IDENTICAL workbook bytes.
    Safe against {145.10}'s per-form dedup: dedup only drops rows and
    recomputes ``sequence``, it never mutates a surviving field's
    ``table_index`` (``xlsx.py``'s ``_dedup_per_form`` docstring — "Preserves
    the original coordinates from the first occurrence").
    """
    from scripts.cocoindex_pipeline.form_extractors.xlsx import _walk_sheet

    wb = openpyxl.load_workbook(
        io.BytesIO(raw_bytes), data_only=True, read_only=False
    )
    sheet_by_table_index: dict[int, str] = {}
    table_index = 0
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        _fields, _next_seq, consumed = _walk_sheet(
            ws, table_index_offset=table_index, sequence_start=0
        )
        for ti in range(table_index, table_index + consumed):
            sheet_by_table_index[ti] = sheet_name
        table_index += consumed
    return sheet_by_table_index


def fill_xlsx_template(
    input_path: str,
    output_path: str,
    field_mappings: list[dict],
) -> dict:
    """Write approved bid responses into an XLSX workbook's cells.

    Args:
        input_path: Path to the base XLSX — the original upload on a first
            pass, or a PRIOR ``template_completions`` output on a re-fill
            pass (BI-22 re-entrancy).
        output_path: Path to save the filled XLSX.
        field_mappings: List of dicts with:
            - table_index (int) — the workbook-wide table counter (see
              ``_xlsx_sheet_by_table_index``).
            - row_index (int), col_index (int) — 1-based openpyxl
              coordinates within the resolved sheet.
            - response_text (str)
            - word_limit (int | None)

    Returns:
        Dict with fields_filled, fields_skipped, fields_failed, truncated, errors.
    """
    with open(input_path, "rb") as f:
        raw_bytes = f.read()

    sheet_by_table_index = _xlsx_sheet_by_table_index(raw_bytes)

    wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), data_only=False, read_only=False)

    filled = 0
    skipped = 0
    failed = 0
    truncated_fields: list[dict] = []
    errors: list[dict] = []

    for mapping in field_mappings:
        table_idx = mapping.get("table_index")
        row_idx = mapping.get("row_index")
        col_idx = mapping.get("col_index")
        response_text = mapping.get("response_text")
        word_limit = mapping.get("word_limit")

        sheet_name = sheet_by_table_index.get(table_idx)
        if sheet_name is None:
            errors.append({
                "table_index": table_idx,
                "row_index": row_idx,
                "error": f"No sheet found for table_index {table_idx}",
            })
            failed += 1
            continue

        if not response_text or not response_text.strip():
            skipped += 1
            continue

        text_to_write, was_truncated = _enforce_word_limit(
            response_text, word_limit
        )
        if was_truncated:
            truncated_fields.append({
                "table_index": table_idx,
                "row_index": row_idx,
                "original_words": len(response_text.split()),
                "limit": word_limit,
            })

        try:
            wb[sheet_name].cell(row=row_idx, column=col_idx).value = text_to_write
            filled += 1
        except Exception as e:
            errors.append({
                "table_index": table_idx,
                "row_index": row_idx,
                "error": str(e),
            })
            failed += 1

    wb.save(output_path)

    return {
        "fields_filled": filled,
        "fields_skipped": skipped,
        "fields_failed": failed,
        "truncated": truncated_fields,
        "errors": errors,
    }


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(
            "Usage: python3 scripts/fill_template.py <template.docx> <output.docx> [mappings.json]",
            file=sys.stderr,
        )
        sys.exit(1)

    template_path = sys.argv[1]
    output_path = sys.argv[2]
    mappings_file = sys.argv[3] if len(sys.argv) > 3 else None

    if mappings_file:
        with open(mappings_file) as f:
            mappings = json.load(f)
    else:
        mappings = json.load(sys.stdin)

    try:
        result = fill_template(template_path, output_path, mappings)
        print(json.dumps(result, indent=2, ensure_ascii=False))
        print(
            f"Filled: {result['fields_filled']}, "
            f"Skipped: {result['fields_skipped']}, "
            f"Failed: {result['fields_failed']}",
            file=sys.stderr,
        )
    except Exception as e:
        print(f"Error filling template: {e}", file=sys.stderr)
        sys.exit(1)
